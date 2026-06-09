"""
Reporte de Eficiencia Vicidial — App standalone sin login
streamlit run vicidial_report.py
"""
import streamlit as st
import pandas as pd
import numpy as np
import plotly.express as px
import plotly.graph_objects as go
import io
from datetime import datetime

st.set_page_config(
    page_title="Reporte Vicidial",
    page_icon="📞",
    layout="wide",
    initial_sidebar_state="collapsed",
)

st.markdown("""
<style>
[data-testid="stAppViewContainer"] { background: #0f1117; }
[data-testid="stSidebar"] { display: none; }
</style>
""", unsafe_allow_html=True)

_VIC_CONTACT  = {"SALE","CALLBK","DEC","NI","XFER","INCALL","CLOSER"}
_VIC_POSITIVE = {"SALE","CALLBK","XFER"}
_VIC_MACHINE  = {"AA","AM","LAMA"}
_VIC_NOANSWER = {"N","NA","NNA","B","BUSY","DC","DROP","QUEUETIMEOUT"}
_VIC_STATUS_LABELS = {
    "SALE":"Venta","CALLBK":"Callback","DEC":"Rechazo","NI":"No interesado",
    "XFER":"Transferido","N":"No contesta","NA":"No contesta","NNA":"No contesta",
    "AA":"Contestador","AM":"Contestador","B":"Ocupado","BUSY":"Ocupado",
    "DC":"Desconectado","DROP":"Caida","LAMA":"Limit.contestador",
    "QUEUETIMEOUT":"Timeout cola","INCALL":"En llamada","DNCL":"No llamar",
}

def _detect_cols(df):
    c = {col.lower().strip(): col for col in df.columns}
    def find(*names):
        for n in names:
            if n in c: return c[n]
        return None
    return {
        "date":     find("call_date","fecha","date","start_time","calldate","call_time"),
        "agent":    find("user","agent","agente","operator","username","agent_user"),
        "status":   find("status","estado","disposition","call_status","result"),
        "duration": find("length_in_sec","duration","duracion","call_duration","seconds","length","talk_time"),
        "campaign": find("campaign_id","campaign","campana","lista","list_id"),
    }

def _load(uploaded):
    name = uploaded.name.lower()
    if name.endswith((".xlsx",".xls")): return pd.read_excel(uploaded)
    if name.endswith(".docx"):
        from docx import Document
        doc = Document(uploaded)
        if doc.tables:
            tbl = doc.tables[0]
            headers = [cell.text.strip() for cell in tbl.rows[0].cells]
            data = [[cell.text.strip() for cell in row.cells] for row in tbl.rows[1:]]
            df = pd.DataFrame(data, columns=headers)
            return df.loc[:, df.columns.str.strip() != ""]
        raise ValueError("No se encontraron tablas en el documento.")
    for enc in ["utf-8","latin-1","cp1252"]:
        for sep in [",",";","\t","|"]:
            try:
                df = pd.read_csv(uploaded, encoding=enc, sep=sep, low_memory=False)
                if len(df.columns) > 1: return df
                uploaded.seek(0)
            except Exception: uploaded.seek(0)
    return pd.read_csv(uploaded, encoding="latin-1", low_memory=False)

def _fmt(s):
    if s is None or (isinstance(s, float) and np.isnan(s)): return "–"
    m, s = divmod(int(s), 60); return f"{m}m {s:02d}s"

def _pct(n, d): return "0%" if d == 0 else f"{n/d*100:.1f}%"

def _kpi(col, label, value, color="#3b82f6"):
    with col:
        st.markdown(f"""
        <div style="background:#1e2535;border-radius:10px;padding:16px 18px;
                    border-left:4px solid {color};margin-bottom:8px">
          <div style="color:#8899aa;font-size:12px;text-transform:uppercase">{label}</div>
          <div style="color:#fff;font-size:26px;font-weight:700">{value}</div>
        </div>""", unsafe_allow_html=True)

def _alert(title, body, color="#dc2626"):
    st.markdown(f"""
    <div style="background:#1c1917;border:1px solid {color};border-radius:10px;
                padding:14px 18px;margin-bottom:8px">
      <div style="color:{color};font-weight:700;font-size:14px">⚠ {title}</div>
      <div style="color:#d1d5db;font-size:13px;margin-top:4px">{body}</div>
    </div>""", unsafe_allow_html=True)

def _critical(total, contacted, machines, no_answer, avg_dur, agent_df):
    pts = []
    cr = contacted/total if total else 0
    if cr < 0.20: pts.append(("rojo","Tasa de contacto CRÍTICA",f"Solo {cr*100:.1f}% contactan (meta 20%). Revisar base y horarios."))
    elif cr < 0.35: pts.append(("naranja","Tasa de contacto baja",f"{cr*100:.1f}% de contacto. Optimizar franjas horarias."))
    if machines/total > 0.30 if total else False: pts.append(("naranja","Alto % contestadores",f"{machines/total*100:.1f}% AM. Activar AMD o ajustar horarios."))
    if avg_dur:
        if avg_dur < 30: pts.append(("rojo","Duración promedio muy corta",f"{avg_dur:.0f}s. Revisar conectividad y scripts."))
        elif avg_dur > 600: pts.append(("naranja","Llamadas muy largas",f"{avg_dur/60:.1f} min promedio. Revisar scripts."))
    if agent_df is not None and "tasa_contacto" in agent_df.columns:
        low = agent_df[agent_df["tasa_contacto"] < 15]
        if len(low): pts.append(("naranja",f"Agentes con contacto < 15%",", ".join(low.index.astype(str)[:5])))
    if not pts: pts.append(("verde","Sin alertas críticas","El reporte no muestra problemas graves."))
    return pts

C = dict(paper_bgcolor="#1e2535", plot_bgcolor="#1e2535",
         font=dict(color="#aaa",size=12), margin=dict(t=36,b=20,l=10,r=10))

# ── APP ───────────────────────────────────────────────────────────────────────
st.markdown("## 📞 Reporte de Eficiencia — Vicidial")
st.markdown("Sube tu archivo de llamadas y obtén el análisis en segundos.")

uploaded = st.file_uploader("Archivo de llamadas (xlsx, csv, txt, docx)", type=None)

if not uploaded:
    with st.expander("¿Qué columnas necesita?"):
        st.markdown("""
| Columna | Nombres reconocidos |
|---------|---------------------|
| Fecha | `call_date`, `fecha`, `start_time` |
| Agente | `user`, `agent`, `agente` |
| Estado | `status`, `disposition`, `estado` |
| Duración (seg) | `length_in_sec`, `duration`, `duracion` |
| Campaña | `campaign_id`, `campaign` |
        """)
    st.stop()

with st.spinner("Analizando..."):
    try:
        df = _load(uploaded)
    except Exception as e:
        st.error(f"Error al leer: {e}"); st.stop()

cols  = _detect_cols(df)
total = len(df)

if cols["status"]:
    df["_st"] = df[cols["status"]].astype(str).str.upper().str.strip()
    contacted = int(df["_st"].isin(_VIC_CONTACT).sum())
    positives = int(df["_st"].isin(_VIC_POSITIVE).sum())
    machines  = int(df["_st"].isin(_VIC_MACHINE).sum())
    no_answer = int(df["_st"].isin(_VIC_NOANSWER).sum())
else:
    contacted = positives = machines = no_answer = 0

if cols["duration"]:
    df["_dur"]  = pd.to_numeric(df[cols["duration"]], errors="coerce").fillna(0)
    avg_dur     = df["_dur"].mean()
    total_min   = df["_dur"].sum() / 60
else:
    avg_dur = total_min = None

if cols["date"]:
    df["_dt"]   = pd.to_datetime(df[cols["date"]], errors="coerce")
    df["_hour"] = df["_dt"].dt.hour

# KPIs
st.markdown("---")
k1,k2,k3,k4,k5 = st.columns(5)
_kpi(k1, "Total llamadas", f"{total:,}")
_kpi(k2, "Tasa de contacto", _pct(contacted,total),
     "#27ae60" if contacted/total>=0.35 else "#f59e0b" if contacted/total>=0.20 else "#e74c3c")
_kpi(k3, "Tasa positiva", _pct(positives,total),
     "#27ae60" if positives/total>=0.15 else "#f59e0b" if positives/total>=0.08 else "#e74c3c")
_kpi(k4, "Duración promedio", _fmt(avg_dur),
     "#27ae60" if avg_dur and 60<=avg_dur<=300 else "#f59e0b")
_kpi(k5, "Minutos totales", f"{total_min:,.0f}" if total_min else "–")

# Puntos críticos
st.markdown("---")
st.markdown("#### ⚠ Puntos Críticos")
agent_df_alerts = None
if cols["agent"] and cols["status"]:
    ag  = pd.concat([df.groupby(cols["agent"]).size().rename("total"),
                     df[df["_st"].isin(_VIC_CONTACT)].groupby(cols["agent"]).size().rename("contactadas")], axis=1).fillna(0)
    ag["tasa_contacto"] = (ag["contactadas"]/ag["total"]*100).round(1)
    agent_df_alerts = ag

pts = _critical(total, contacted, machines, no_answer, avg_dur, agent_df_alerts)
cm  = {"rojo":"#dc2626","naranja":"#f59e0b","verde":"#22c55e"}
for sev, title, body in pts:
    if sev == "verde": st.success(f"✅ **{title}** — {body}")
    else: _alert(title, body, cm[sev])

# Gráficos
st.markdown("---")
t1,t2,t3,t4 = st.tabs(["📊 Distribución","⏰ Por Hora","👤 Por Agente","📈 Tendencia"])

with t1:
    ca,cb = st.columns(2)
    with ca:
        if cols["status"]:
            sc = df["_st"].value_counts().head(12).reset_index()
            sc.columns = ["Status","Llamadas"]
            sc["Label"] = sc["Status"].map(_VIC_STATUS_LABELS).fillna(sc["Status"])
            fig = px.pie(sc, values="Llamadas", names="Label", title="Por Estado", hole=0.4,
                         color_discrete_sequence=px.colors.qualitative.Set3)
            fig.update_layout(**C, height=320); st.plotly_chart(fig, use_container_width=True)
    with cb:
        if cols["status"]:
            cats = {"Contacto":contacted,"No contesta":no_answer,"Contestador AM":machines,
                    "Otros":max(0,total-contacted-no_answer-machines)}
            fig2 = px.bar(x=list(cats.keys()), y=list(cats.values()), title="Resultado",
                          color=list(cats.keys()),
                          color_discrete_map={"Contacto":"#22c55e","No contesta":"#ef4444",
                                              "Contestador AM":"#f59e0b","Otros":"#6b7280"})
            fig2.update_layout(**C, height=320, showlegend=False); st.plotly_chart(fig2, use_container_width=True)
    if cols["duration"]:
        d = df["_dur"][df["_dur"].between(1,3600)]
        fig3 = px.histogram(d, nbins=50, title="Duración (seg)", color_discrete_sequence=["#3b82f6"])
        fig3.add_vline(x=d.mean(), line_dash="dash", line_color="#f59e0b", annotation_text=f"Prom: {d.mean():.0f}s")
        fig3.update_layout(**C); st.plotly_chart(fig3, use_container_width=True)

with t2:
    if cols["date"] and "_hour" in df.columns:
        hc = df.groupby("_hour").size().reindex(range(24), fill_value=0)
        fh = go.Figure()
        fh.add_trace(go.Bar(x=hc.index, y=hc.values, name="Llamadas", marker_color="#3b82f6", opacity=0.7))
        if cols["status"]:
            hcon = df[df["_st"].isin(_VIC_CONTACT)].groupby("_hour").size().reindex(range(24),fill_value=0)
            fh.add_trace(go.Scatter(x=hcon.index, y=hcon.values, name="Contactos",
                                    line=dict(color="#22c55e",width=2), mode="lines+markers"))
            st.info(f"Mejor hora: **{hcon.idxmax()}:00** ({hcon.max():,} contactos)")
        fh.update_layout(**C, title="Llamadas por Hora",
                         xaxis=dict(title="Hora",tickmode="linear",dtick=1),
                         yaxis=dict(title="Llamadas"), legend=dict(x=0,y=1.1,orientation="h"))
        st.plotly_chart(fh, use_container_width=True)
        dow_o = ["Monday","Tuesday","Wednesday","Thursday","Friday","Saturday","Sunday"]
        dow_e = ["Lunes","Martes","Miércoles","Jueves","Viernes","Sábado","Domingo"]
        dow = df.groupby(df["_dt"].dt.day_name()).size().reindex([d for d in dow_o if d in df["_dt"].dt.day_name().values],fill_value=0)
        dow.index = [dow_e[dow_o.index(d)] for d in dow.index]
        fd = px.bar(x=dow.index, y=dow.values, title="Por Día de Semana", color_discrete_sequence=["#8b5cf6"])
        fd.update_layout(**C); st.plotly_chart(fd, use_container_width=True)
    else: st.info("No se detectó columna de fecha.")

with t3:
    if cols["agent"]:
        res = {"total": df.groupby(cols["agent"]).size().rename("total")}
        if cols["status"]:
            res["contactadas"] = df[df["_st"].isin(_VIC_CONTACT)].groupby(cols["agent"]).size()
        if cols["duration"]:
            res["avg_seg"] = df.groupby(cols["agent"])["_dur"].mean().round(0)
        adf = pd.DataFrame(res).fillna(0).sort_values("total",ascending=False)
        if "contactadas" in adf.columns:
            adf["contacto_%"] = (adf["contactadas"]/adf["total"]*100).round(1)
        fa = px.bar(adf.head(20).reset_index(), x=cols["agent"], y="total",
                    title="Top Agentes", color_discrete_sequence=["#3b82f6"])
        fa.update_layout(**C); st.plotly_chart(fa, use_container_width=True)
        if "contacto_%" in adf.columns:
            fac = px.bar(adf.head(20).sort_values("contacto_%",ascending=False).reset_index(),
                         x=cols["agent"], y="contacto_%", title="Tasa de Contacto (%)",
                         color="contacto_%", color_continuous_scale=["#ef4444","#f59e0b","#22c55e"],range_color=[0,60])
            fac.add_hline(y=35, line_dash="dash", line_color="#22c55e", annotation_text="Meta 35%")
            fac.update_layout(**C); st.plotly_chart(fac, use_container_width=True)
        st.dataframe(adf.reset_index().head(30), use_container_width=True, hide_index=True)
    else: st.info("No se detectó columna de agente.")

with t4:
    if cols["date"] and "_dt" in df.columns:
        df2 = df.copy(); df2["_date"] = df2["_dt"].dt.date
        daily = df2.groupby("_date").size().reset_index(name="llamadas").dropna()
        if len(daily) > 1:
            ft = go.Figure()
            ft.add_trace(go.Scatter(x=daily["_date"], y=daily["llamadas"], mode="lines+markers",
                                    name="Llamadas/día", line=dict(color="#3b82f6",width=2),
                                    fill="tozeroy", fillcolor="rgba(59,130,246,0.1)"))
            if cols["status"]:
                dc = df2[df2["_st"].isin(_VIC_CONTACT)].groupby("_date").size().reset_index(name="contactos")
                dm = daily.merge(dc,on="_date",how="left").fillna(0)
                dm["tasa"] = dm["contactos"]/dm["llamadas"]*100
                ft.add_trace(go.Scatter(x=dm["_date"],y=dm["tasa"],name="Tasa contacto %",yaxis="y2",
                                        line=dict(color="#22c55e",width=2,dash="dot")))
                ft.update_layout(yaxis2=dict(title="Tasa %",overlaying="y",side="right",color="#22c55e"))
            ft.update_layout(**C, title="Volumen Diario",
                             xaxis=dict(title="Fecha"),yaxis=dict(title="Llamadas"))
            st.plotly_chart(ft, use_container_width=True)
        else: st.info("Solo un día de datos.")
    else: st.info("No se detectó columna de fecha.")

# Exportar
st.markdown("---")
buf = io.BytesIO()
with pd.ExcelWriter(buf, engine="openpyxl") as w:
    pd.DataFrame({"Métrica":["Total llamadas","Contactados","Tasa contacto","Positivos",
                              "Tasa positiva","Contestadores AM","No contesta","Duración prom.","Min. totales"],
                  "Valor":[total,contacted,_pct(contacted,total),positives,_pct(positives,total),
                           machines,no_answer,_fmt(avg_dur),f"{total_min:.0f}" if total_min else "–"]}
                 ).to_excel(w, sheet_name="Resumen", index=False)
    pd.DataFrame([(s,t,b) for s,t,b in pts],columns=["Severidad","Título","Descripción"]
                ).to_excel(w, sheet_name="Puntos Críticos", index=False)
    if cols["status"]: df["_st"].value_counts().reset_index().to_excel(w, sheet_name="Por Estado", index=False)
    if agent_df_alerts is not None: agent_df_alerts.reset_index().to_excel(w, sheet_name="Por Agente", index=False)
buf.seek(0)
st.download_button("⬇️ Descargar Reporte Excel", data=buf,
    file_name=f"reporte_vicidial_{datetime.now().strftime('%Y%m%d_%H%M')}.xlsx",
    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", type="primary")
