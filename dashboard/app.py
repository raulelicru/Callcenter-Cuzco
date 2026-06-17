"""
Dashboard de Cobranza - Sistema Integral
streamlit run dashboard/app.py
"""

import sys, os
sys.path.insert(0, os.path.join(os.path.dirname(__file__), "..", "src"))

import streamlit as st
import pandas as pd
import numpy as np
import plotly.express as px
import plotly.graph_objects as go
import joblib, io, re, unicodedata
from pathlib import Path
from datetime import datetime, date

# ── Constantes Vicidial ───────────────────────────────────────────────────────
_VIC_MACHINE  = {"AA", "AM", "LAMA"}
_VIC_NOANSWER = {"N", "NA", "NNA", "B", "BUSY", "DC", "DROP", "QUEUETIMEOUT"}
_VIC_STATUS_LABELS = {
    "SALE":"Venta","CALLBK":"Callback","DEC":"Rechazo","NI":"No interesado",
    "XFER":"Transferido","N":"No contesta","NA":"No contesta","NNA":"No contesta",
    "AA":"Contestador","AM":"Contestador","B":"Ocupado","BUSY":"Ocupado",
    "DC":"Desconectado","DROP":"Caida","LAMA":"Limit.contestador",
    "QUEUETIMEOUT":"Timeout cola","INCALL":"En llamada","DNCL":"No llamar",
}

# Reglas de la campaña GRAL (reporte diario de 6 archivos)
_VIC_GRAL_PROMESA   = {"04", "21"}
_VIC_GRAL_SALUDO    = "01"
_VIC_GRAL_HUMANOS   = {"01", "02", "04", "09", "14", "18", "19", "21"}
_VIC_GRAL_LABELS    = {
    "01": "Cuelga en saludo (rechazo temprano)",
    "02": "Contacto - no interesado",
    "04": "Promesa de pago",
    "09": "Contacto humano",
    "14": "Contacto humano",
    "18": "Contacto humano",
    "19": "Contacto humano",
    "21": "Promesa de pago (alterna)",
}

# Reglas de la campaña Coquimbo (mismos 3 archivos, otra tipificación interna)
_COQ_PROMESA = {"1B", "1O"}
_COQ_SALUDO  = "1L"
_COQ_HUMANOS = {"02", "04", "06", "1B", "1C", "1D", "1E", "1F", "1G", "1H", "1I", "1J", "1K",
                "1L", "1N", "1O", "2B", "2C", "2D", "2E", "2F", "3C", "3D", "3E"}
_COQ_LABELS  = {
    "02":     "Acuerdo autorizado",
    "04":     "Entrante no relacionada",
    "06":     "Trasferencia de llamada",
    "1B":     "Promesa de pago",
    "1C":     "Informa ya pagó",
    "1D":     "Consultor no puede pagar",
    "1E":     "No se inscribió / falsificación",
    "1F":     "No pasó pedido",
    "1G":     "Presta código a empresaria",
    "1H":     "Entregó dinero a empresaria",
    "1I":     "Ajuste no procesado",
    "1J":     "Familiar usa código de consultor",
    "1K":     "Canceló por error a otro código",
    "1L":     "Cuelga en saludo (rechazo temprano)",
    "1N":     "Posterga pago",
    "1O":     "Promesa de pago (alterna)",
    "2B":     "Recado familiar",
    "2C":     "Recado trabajo",
    "2D":     "Recado amigo",
    "2E":     "Recado vecino",
    "2F":     "Mensaje a terceros",
    "3B":     "Línea incorrecta, no enlaza",
    "3C":     "No contesta",
    "3D":     "Número no corresponde",
    "3E":     "Ya no vive / trabaja ahí",
    "3H":     "Buzón celular",
    "AA":     "Buzón de voz o contestadora automática",
    "AB":     "Agente no disponible",
    "DROP":   "Llamada descartada o colgada por el sistema antes de conectarse con un agente",
    "INCALL": "Llamada exitosa",
    "NA":     "No contesta",
    "PDROP":  "Llamada colgada antes de atención",
}

from database import (
    init_db, get_clientes_by_ids, upsert_clientes_batch,
    log_carga, get_cargas_historico, get_metricas_globales, get_all_clientes_df,
    get_all_empresas, create_empresa,
)
from auth import authenticate, create_user, get_all_users, toggle_user_status, update_password

# ── Rutas absolutas ───────────────────────────────────────────────────────────
BASE_DIR   = Path(__file__).parent.parent
MODELS_DIR = BASE_DIR / "models"
DATA_DIR   = BASE_DIR / "data"

# ── Config de página ──────────────────────────────────────────────────────────
st.set_page_config(
    page_title="Cobranza | Call Center Cuzco",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ── Tema (Oscuro / Claro) ─────────────────────────────────────────────────────
st.session_state.setdefault("theme", "Oscuro")

_LIGHT_CSS = """
.stApp {
    --background-color: #ffffff;
    --secondary-background-color: #f1f5f9;
    --text-color: #0f172a;
    background-color: #ffffff !important;
}
[data-testid="stSidebar"] {
    background: linear-gradient(180deg, #f8fafc 0%, #eef2f7 100%) !important;
    border-right: 1px solid #e2e8f0 !important;
}
[data-testid="stSidebar"] .stButton button { color: #475569 !important; }
[data-testid="stSidebar"] .stButton button:hover { background: #e2e8f0 !important; color: #0f172a !important; }
hr { border-color: #e2e8f0 !important; }
[data-testid="stDataFrame"] { border: 1px solid #e2e8f0 !important; }

/* Texto general legible sobre fondo blanco (no afecta tarjetas con fondo propio) */
.stApp h1, .stApp h2, .stApp h3, .stApp h4, .stApp h5, .stApp h6,
.stApp p, .stApp li, .stApp label,
[data-testid="stCaptionContainer"],
[data-testid="stMetricLabel"], [data-testid="stMetricValue"],
[data-testid="stExpander"] summary,
[data-testid="stExpander"] summary span,
[data-testid="stExpander"] p,
[data-testid="stMarkdownContainer"] {
    color: #0f172a !important;
}
[data-testid="stSidebar"] [data-testid="stCaptionContainer"] { color: #64748b !important; }

/* Cargador de archivos (drag & drop) legible en modo claro */
[data-testid="stFileUploaderDropzone"] {
    background: #f1f5f9 !important;
    border: 1px dashed #cbd5e1 !important;
}
[data-testid="stFileUploaderDropzone"] * { color: #0f172a !important; }
[data-testid="stFileUploaderDropzone"] button {
    background: #ffffff !important;
    border: 1px solid #cbd5e1 !important;
    color: #0f172a !important;
}
[data-testid="stFileUploaderFile"] { color: #0f172a !important; }
"""

# ── CSS Global ────────────────────────────────────────────────────────────────
_BASE_CSS = """
/* Fondo general */
.stApp { background-color: #0f1117; }

/* Sidebar */
[data-testid="stSidebar"] {
    background: linear-gradient(180deg, #161b27 0%, #1a2035 100%);
    border-right: 1px solid #2a3045;
}

/* Tarjeta métrica personalizada */
.metrica {
    background: #1e2535;
    border-radius: 10px;
    padding: 18px 20px;
    border-left: 4px solid;
    margin-bottom: 8px;
}
.metrica .metrica-valor { font-size: 2rem; font-weight: 700; color: #ffffff !important; margin: 0; }
.metrica .metrica-label { font-size: 0.78rem; color: #8899aa !important; margin: 0; text-transform: uppercase; letter-spacing: 0.05em; }
.metrica .metrica-delta { font-size: 0.85rem; margin-top: 4px; }

/* Badge de segmento */
.badge {
    display: inline-block;
    padding: 3px 10px;
    border-radius: 20px;
    font-size: 0.75rem;
    font-weight: 700;
    letter-spacing: 0.05em;
}
.badge-alto  { background: #1a3a1a; color: #4ade80 !important; border: 1px solid #27ae60; }
.badge-medio { background: #3a2a00; color: #fbbf24 !important; border: 1px solid #f39c12; }
.badge-bajo  { background: #3a0f0f; color: #f87171 !important; border: 1px solid #e74c3c; }

/* Card de estrategia */
.card-estrategia {
    background: #1e2535;
    border-radius: 12px;
    padding: 20px 24px;
    margin-bottom: 16px;
}

/* Login card */
.login-card {
    background: #1e2535;
    border-radius: 16px;
    padding: 40px 36px;
    border: 1px solid #2a3045;
    box-shadow: 0 8px 32px rgba(0,0,0,0.4);
}
.login-card .login-title { font-size: 1.8rem; font-weight: 800; color: #fff !important; margin-bottom: 4px; }
.login-card .login-sub   { color: #6b7a99 !important; font-size: 0.95rem; margin-bottom: 28px; }

/* Botones de navegacion */
[data-testid="stSidebar"] .stButton button {
    width: 100%;
    text-align: left;
    background: transparent;
    border: none;
    color: #8899aa;
    padding: 8px 12px;
    border-radius: 8px;
    font-size: 0.9rem;
    margin-bottom: 2px;
    transition: all 0.15s;
}
[data-testid="stSidebar"] .stButton button:hover {
    background: #2a3045;
    color: #ffffff;
}

/* Divider */
hr { border-color: #2a3045; }

/* Dataframe */
[data-testid="stDataFrame"] { border: 1px solid #2a3045; border-radius: 8px; }

/* KPI highlight bar */
.kpi-bar {
    display: flex; gap: 12px; flex-wrap: wrap; margin-bottom: 20px;
}
.kpi-item {
    flex: 1; min-width: 140px;
    background: #1e2535;
    border-radius: 10px;
    padding: 16px;
    border-top: 3px solid;
    text-align: center;
}
"""

_css = _BASE_CSS + (_LIGHT_CSS if st.session_state["theme"] == "Claro" else "")
st.markdown(f"<style>{_css}</style>", unsafe_allow_html=True)

# ── Colores por segmento ──────────────────────────────────────────────────────
COLORES = {"ALTO": "#27ae60", "MEDIO": "#f39c12", "BAJO": "#e74c3c"}

# ── Estrategias ───────────────────────────────────────────────────────────────
ESTRATEGIAS = {
    "ALTO": {
        "canal":      "WhatsApp Business / SMS / Llamada IVR Autopago",
        "accion":     "Mensaje WhatsApp/SMS automatizado con link de pago + IVR de recordatorio",
        "oferta":     "Pago total con descuento / 2 cuotas sin interes / Autopago inmediato",
        "frecuencia": "1 WhatsApp + 1 SMS por semana | Max. 1 llamada IVR semanal",
        "escalacion": "Sin respuesta en 7 dias -> Pasar a gestion MEDIO con agente",
        "script":     "WhatsApp: Hola [Nombre], te recordamos que tienes un saldo pendiente de S/[monto]. Regularizalo hoy y evita cargos adicionales. Paga aqui: [link]",
        "kpis":       "Tasa de respuesta WhatsApp >= 25% | Conversion >= 20% | Costo S/ 0.20-0.50",
        "referencia": "Estrategia Digital First — marcador predictivo en standby",
    },
    "MEDIO": {
        "canal":      "Marcador Predictivo + Agente Humano (AMD activado) + WhatsApp de seguimiento",
        "accion":     "Llamada outbound de negociacion activa + WhatsApp post-llamada con resumen de acuerdo",
        "oferta":     "Plan 3-6 cuotas / Condonacion de intereses moratorios / Refinanciamiento",
        "frecuencia": "Max. 3 llamadas/dia | Mejor hora: 9-11am y 6-8pm L-V | WhatsApp de seguimiento mismo dia",
        "escalacion": "2 Promesas de Pago rotas -> Escalar a gestion BAJO intensiva",
        "script":     "Buenos dias [Nombre], le llamo del Call Center Cuzco por su cuenta. Entiendo que puede tener dificultades, por eso le ofrezco regularizar en [N] cuotas. ¿Le parece bien comenzar hoy?",
        "kpis":       "RPC >= 45% | PTP >= 30% | Promesas cumplidas >= 65% | Costo S/ 2.00-3.50",
        "referencia": "Script ACED: Acknowledge + Create urgency + Empathize + Deal",
    },
    "BAJO": {
        "canal":      "Agente Senior Especializado / Llamada Directa / SMS y WhatsApp de presion",
        "accion":     "Llamada directa de agente senior + SMS/WhatsApp con aviso formal de mora + oferta de settlement",
        "oferta":     "Score 22-33: Refinanciamiento largo plazo | Score 11-21: Quita 20-40% del saldo | Score 1-10: Settlement minimo + cierre de cuenta",
        "frecuencia": "1 llamada diaria de agente senior | SMS/WhatsApp de seguimiento a las 2 horas post-llamada",
        "escalacion": "DPD > 180 y sin acuerdo -> Notificacion formal por llamada + SMS de pre-corte",
        "script":     "Sr./Sra. [Nombre], le llama [Agente] del area de Recuperaciones del Call Center Cuzco. Su cuenta tiene [DPD] dias de mora por S/[monto]. Tenemos una oferta especial de settlement disponible solo por hoy. ¿Puede atendernos?",
        "kpis":       "Contacto efectivo >= 30% | Settlement >= 15% | Recuperacion >= 8% del saldo",
        "referencia": "Gestion intensiva telefonica — sin derivacion externa",
    },
}

# ─────────────────────────────────────────────────────────────────────────────
# HELPERS
# ─────────────────────────────────────────────────────────────────────────────

@st.cache_resource
def load_model():
    p = MODELS_DIR / "pipeline_random_forest.pkl"
    return joblib.load(p) if p.exists() else None


def _generate_plan(row) -> str:
    """Genera un plan de accion personalizado por cliente segun su perfil."""
    seg     = str(row.get("segmento", "MEDIO"))
    dpd     = int(row.get("dpd", 0) or 0)
    saldo   = float(row.get("saldo_total", 0) or 0)
    rpc     = float(row.get("rpc_rate", 0) or 0)
    estado  = str(row.get("ultimo_estado_marcado", "") or "")
    p_rotas = int(row.get("promesas_rotas", 0) or 0)

    if seg == "ALTO":
        if rpc >= 0.55:
            return (f"[WhatsApp] Enviar hoy: 'Hola [Nombre], tienes S/{saldo:,.0f} pendiente. "
                    f"Pagalo hoy y evitamos mas cargos. Ingresa aqui: [link_pago]'")
        elif rpc >= 0.25:
            return f"[SMS] Automatico: 'Deuda S/{saldo:,.0f}. Paga en [link]. Consultas: [tel]'"
        else:
            return f"[IVR] Llamada automatica de recordatorio: saldo S/{saldo:,.0f}"

    elif seg == "MEDIO":
        if p_rotas >= 2:
            return (f"[Llamada urgente] {p_rotas} PTPs rotas. Agente negocia nueva fecha "
                    f"con compromiso formal. Saldo S/{saldo:,.0f}")
        elif "PROMESA" in estado.upper():
            return (f"[WhatsApp] Seguimiento PTP: 'Recordamos su compromiso de pago "
                    f"S/{saldo:,.0f}. Confirme si cumplira en la fecha acordada.'")
        else:
            cuota = saldo / 3
            return (f"[Marcador + Agente] Ofrecer 3 cuotas de S/{cuota:,.0f} "
                    f"con condonacion de intereses. DPD={dpd} dias.")

    else:  # BAJO
        if dpd > 150:
            desc = 40
        elif dpd > 90:
            desc = 25
        elif dpd > 60:
            desc = 15
        else:
            desc = 10
        settlement = saldo * (1 - desc / 100)
        return (f"[Agente Senior] Settlement S/{settlement:,.0f} (quita {desc}% "
                f"sobre S/{saldo:,.0f}). Llamada directa + SMS seguimiento. DPD={dpd} dias.")


def score_df(df_raw, pipeline):
    from model import score_portfolio
    df_s = score_portfolio(df_raw, pipeline)
    for k in ["canal", "accion", "oferta", "frecuencia"]:
        df_s[f"estrategia_{k}"] = df_s["segmento"].map(
            lambda s, k=k: ESTRATEGIAS.get(s, {}).get(k, "")
        )
    return df_s


def process_upload(df_raw, pipeline, usuario, filename):
    if "cliente_id" not in df_raw.columns:
        posibles = [c for c in df_raw.columns if "id" in c.lower() or "cliente" in c.lower()]
        if posibles:
            df_raw = df_raw.rename(columns={posibles[0]: "cliente_id"})
        else:
            df_raw.insert(0, "cliente_id", [f"CLI-{i:06d}" for i in range(1, len(df_raw)+1)])

    # ── Deduplicacion: eliminar registros repetidos por cliente_id ──────────────
    df_raw["cliente_id"] = df_raw["cliente_id"].astype(str)
    n_original   = len(df_raw)
    df_raw       = df_raw.drop_duplicates(subset=["cliente_id"], keep="last").reset_index(drop=True)
    n_duplicados = n_original - len(df_raw)

    # ── Detectar conocidos vs nuevos en BD ──────────────────────────────────────
    ids = df_raw["cliente_id"].tolist()
    df_ex = get_clientes_by_ids(ids, _eid())
    ids_conocidos = set(df_ex["cliente_id"].tolist()) if len(df_ex) > 0 else set()
    n_conocidos = sum(1 for i in ids if i in ids_conocidos)
    n_nuevos    = len(ids) - n_conocidos

    # ── Scoring + estrategia + plan personalizado ───────────────────────────────
    df_scored = score_df(df_raw, pipeline)
    df_scored["es_nuevo"]         = ~df_scored["cliente_id"].isin(ids_conocidos)
    df_scored["estado_carga"]     = df_scored["es_nuevo"].map({True: "Nuevo", False: "Actualizado"})
    df_scored["plan_personalizado"] = df_scored.apply(_generate_plan, axis=1)

    # ── Guardar en BD ───────────────────────────────────────────────────────────
    carga_id = log_carga(usuario, filename, len(df_scored), n_nuevos, n_conocidos, _eid())
    upsert_clientes_batch(df_scored, carga_id, _eid())

    return df_scored, {
        "total":        len(df_scored),
        "nuevos":       n_nuevos,
        "conocidos":    n_conocidos,
        "duplicados":   n_duplicados,
        "original":     n_original,
    }


def export_excel(df):
    buf = io.BytesIO()
    plan_cols = [c for c in [
        "cliente_id", "score_operativo", "segmento", "plan_personalizado",
        "dpd", "saldo_total", "prob_pago", "rpc_rate", "bucket_mora",
        "estrategia_canal", "estrategia_oferta", "estado_carga",
    ] if c in df.columns]

    with pd.ExcelWriter(buf, engine="openpyxl") as w:
        # Hoja 1: Plan personalizado por cliente (la mas importante)
        df[plan_cols].sort_values("score_operativo", ascending=False).to_excel(
            w, index=False, sheet_name="Plan por Cliente"
        )
        # Hoja 2-4: Por segmento con plan
        for seg in ["ALTO", "MEDIO", "BAJO"]:
            sub = df[df["segmento"] == seg][plan_cols]
            if len(sub) > 0:
                sub.to_excel(w, index=False, sheet_name=f"Segmento {seg}")
        # Hoja 5: Resumen ejecutivo
        resumen = (
            df.groupby("segmento")
            .agg(clientes=("cliente_id","count"), score_prom=("score_operativo","mean"),
                 prob_prom=("prob_pago","mean"), saldo=("saldo_total","sum"))
            .round(2)
        )
        resumen["pct_cartera"] = (resumen["clientes"] / len(df) * 100).round(1)
        resumen.to_excel(w, sheet_name="Resumen Ejecutivo")
        # Hoja 6: Guia de estrategias
        guia = pd.DataFrame([
            {"Segmento": s, "Canal": ESTRATEGIAS[s]["canal"],
             "Accion": ESTRATEGIAS[s]["accion"], "Oferta": ESTRATEGIAS[s]["oferta"],
             "KPIs": ESTRATEGIAS[s]["kpis"]}
            for s in ["ALTO", "MEDIO", "BAJO"]
        ])
        guia.to_excel(w, index=False, sheet_name="Guia Estrategias")
    return buf.getvalue()


def card_metrica(label, valor, delta=None, color="#3b82f6"):
    delta_html = f'<p class="metrica-delta" style="color:{color}">{delta}</p>' if delta else ""
    st.markdown(f"""
    <div class="metrica" style="border-left-color:{color}">
        <p class="metrica-label">{label}</p>
        <p class="metrica-valor">{valor}</p>
        {delta_html}
    </div>""", unsafe_allow_html=True)


# ─────────────────────────────────────────────────────────────────────────────
# LOGIN
# ─────────────────────────────────────────────────────────────────────────────

def _eid() -> int:
    """Retorna el empresa_id del usuario en sesión."""
    return st.session_state.get("user", {}).get("empresa_id", 1)


def show_login():
    _, col, _ = st.columns([1, 1.1, 1])
    with col:
        st.markdown("<br><br><br>", unsafe_allow_html=True)
        st.markdown("""
        <div class="login-card">
            <p class="login-title">Call Center Cuzco</p>
            <p class="login-sub">Sistema Predictivo de Cobranza</p>
        </div>
        """, unsafe_allow_html=True)

        st.markdown("<br>", unsafe_allow_html=True)

        # Selector de empresa
        try:
            empresas_df = get_all_empresas()
            if len(empresas_df) > 0:
                empresa_opciones = dict(zip(empresas_df["nombre"], empresas_df["id"]))
            else:
                st.warning("No se encontraron empresas en la base de datos.")
                empresa_opciones = {"Cuzco": 1, "Coquimbo": 2}
        except Exception as e:
            st.error(f"Error al cargar empresas: {e}")
            empresa_opciones = {"Cuzco": 1, "Coquimbo": 2}

        with st.form("login_form"):
            empresa_sel = st.selectbox("Empresa", list(empresa_opciones.keys()))
            usuario  = st.text_input("Usuario", placeholder="Ingresa tu usuario")
            password = st.text_input("Contrasena", type="password", placeholder="••••••••")
            ingresar = st.form_submit_button("Ingresar", use_container_width=True, type="primary")

        if ingresar:
            if not usuario or not password:
                st.error("Completa usuario y contrasena.")
            else:
                empresa_id = empresa_opciones.get(empresa_sel, 1)
                try:
                    user = authenticate(usuario, password, empresa_id)
                except Exception as e:
                    st.error(f"Error de conexión: {e}")
                    st.stop()
                if user:
                    st.session_state["user"] = user
                    st.session_state["empresa_nombre"] = empresa_sel
                    st.session_state["page"] = "vicidial"
                    st.rerun()
                else:
                    st.error("Usuario o contrasena incorrectos.")

        st.markdown("<br>", unsafe_allow_html=True)
        st.caption("Acceso restringido. Contacta al Administrador si tienes problemas.")


# ─────────────────────────────────────────────────────────────────────────────
# SIDEBAR
# ─────────────────────────────────────────────────────────────────────────────

def show_sidebar():
    user = st.session_state["user"]
    with st.sidebar:
        empresa_nombre = st.session_state.get("empresa_nombre", "")
        st.markdown(f"""
        <div style="padding: 16px 8px 8px;">
            <div style="font-size:0.75rem;color:#3b82f6;font-weight:600;text-transform:uppercase;letter-spacing:0.05em">{empresa_nombre}</div>
            <div style="font-size:1.1rem;font-weight:700;color:#fff">{user['nombre']}</div>
            <div style="font-size:0.8rem;color:#6b7a99;margin-top:2px">
                {'Administrador' if user['rol']=='admin' else 'Colaborador'}
            </div>
        </div>
        """, unsafe_allow_html=True)
        st.divider()

        pages = {
            "inicio":      ("  Inicio", "Resumen general: total de clientes, score promedio, "
                             "probabilidad de pago y saldo total de la cartera, por segmento."),
            "cargar":      ("  Cargar Cartera", "Sube el archivo de tu cartera de clientes para "
                             "calcular su score, segmento y plan de cobranza personalizado."),
            "analisis":    ("  Analisis", "Gráficos de la cartera: score vs. mora, buckets de mora, "
                             "RPC por segmento, último estado de marcado y saldos."),
            "historial":   ("  Historial", "Historial de cargas realizadas a la base de datos "
                             "(fecha, usuario, registros nuevos y actualizados)."),
            "estrategias": ("  Estrategias", "Estrategias de cobranza recomendadas por segmento "
                             "(canal, oferta, frecuencia, script y KPIs)."),
            "vicidial":    ("  Reporte Coquimbo", "Genera los 3 reportes diarios de la campaña Coquimbo "
                             "(Contactabilidad, Recontacto y Tipificación) a partir de los 3 archivos del día.")
                           if st.session_state.get("empresa_nombre") == "Coquimbo" else
                           ("  Reporte Vicidial", "Genera los 3 reportes diarios de la campaña GRAL "
                             "(Contactabilidad, Recontacto y Tipificación) a partir de los 6 archivos del día."),
            "campana":     ("  Reporte Campaña", "Reporte diario de campaña: resumen de disposiciones, "
                             "rendimiento por ejecutivo, promesas de pago, alertas automáticas y plan de acción."),
        }
        if user["rol"] == "admin":
            pages["admin"] = ("  Panel Admin", "Gestión de usuarios (crear, activar/desactivar, "
                               "cambiar contraseña) y, para la empresa principal, alta de nuevas empresas.")

        cur = st.session_state.get("page", "inicio")
        for key, (label, desc) in pages.items():
            is_active = cur == key
            btn_style = "primary" if is_active else "secondary"
            if st.button(label, use_container_width=True, type=btn_style, key=f"nav_{key}", help=desc):
                st.session_state["page"] = key
                st.rerun()

        st.divider()
        if st.button("Cerrar Sesion", use_container_width=True, key="logout"):
            st.session_state.clear()
            st.rerun()

        st.divider()
        tema = st.radio("Apariencia", ["Oscuro", "Claro"],
                         index=0 if st.session_state["theme"] == "Oscuro" else 1,
                         horizontal=True, key="theme_radio")
        if tema != st.session_state["theme"]:
            st.session_state["theme"] = tema
            st.rerun()

        st.markdown("<br>", unsafe_allow_html=True)
        st.caption("v2.0 — Call Center Cuzco")


# ─────────────────────────────────────────────────────────────────────────────
# PAGINA: INICIO
# ─────────────────────────────────────────────────────────────────────────────

def page_inicio():
    st.markdown("## Resumen General de la Cartera")
    st.markdown(f"*Actualizado: {datetime.now().strftime('%d/%m/%Y %H:%M')}*")
    st.divider()

    m = get_metricas_globales(_eid())
    total = m["total_clientes"]

    if total == 0:
        st.info("La base de datos esta vacia. Ve a **Cargar Cartera** para comenzar.")
        return

    seg   = m["por_segmento"]
    alto  = seg.get("ALTO",  {}).get("count", 0)
    medio = seg.get("MEDIO", {}).get("count", 0)
    bajo  = seg.get("BAJO",  {}).get("count", 0)

    # KPIs
    c1, c2, c3, c4, c5 = st.columns(5)
    with c1: card_metrica("Total en Base de Datos", f"{total:,}", color="#3b82f6")
    with c2: card_metrica("Score Promedio", f"{m['avg_score']}", color="#8b5cf6")
    with c3: card_metrica("Segmento ALTO", f"{alto:,}", f"{alto/total*100:.1f}% de la cartera", "#27ae60")
    with c4: card_metrica("Segmento MEDIO", f"{medio:,}", f"{medio/total*100:.1f}% de la cartera", "#f39c12")
    with c5: card_metrica("Segmento BAJO", f"{bajo:,}", f"{bajo/total*100:.1f}% de la cartera", "#e74c3c")

    st.markdown("<br>", unsafe_allow_html=True)

    df_db = get_all_clientes_df(limit=10000, empresa_id=_eid())
    col_l, col_r = st.columns(2)

    with col_l:
        st.markdown("#### Composicion por Segmento")
        data_pie = {
            "Segmento": list(seg.keys()),
            "Clientes": [v["count"] for v in seg.values()],
        }
        fig = px.pie(data_pie, values="Clientes", names="Segmento",
                     color="Segmento", color_discrete_map=COLORES,
                     hole=0.5, template="plotly_dark")
        fig.update_layout(
            height=340,
            paper_bgcolor="#1e2535", plot_bgcolor="#1e2535",
            font_color="#aaa",
            legend=dict(orientation="h", y=-0.1),
            margin=dict(t=20, b=20),
        )
        fig.update_traces(textinfo="percent+label")
        st.plotly_chart(fig, use_container_width=True)

    with col_r:
        st.markdown("#### Distribucion de Score (1-100)")
        if len(df_db) > 0:
            fig2 = px.histogram(
                df_db, x="score_operativo", nbins=40,
                color="segmento", color_discrete_map=COLORES,
                template="plotly_dark",
                labels={"score_operativo": "Score Operativo", "count": "Clientes"},
            )
            fig2.update_layout(
                height=340,
                paper_bgcolor="#1e2535", plot_bgcolor="#1e2535",
                font_color="#aaa", bargap=0.05,
                margin=dict(t=20, b=20),
                legend_title_text="Segmento",
            )
            st.plotly_chart(fig2, use_container_width=True)

    # Saldo por segmento
    st.markdown("#### Saldo Total por Segmento (S/)")
    saldo_data = [
        {"Segmento": k, "Saldo": round(v.get("saldo", 0), 0)}
        for k, v in seg.items()
    ]
    if saldo_data:
        fig3 = px.bar(
            pd.DataFrame(saldo_data), x="Segmento", y="Saldo",
            color="Segmento", color_discrete_map=COLORES,
            text_auto=".3s", template="plotly_dark",
        )
        fig3.update_layout(
            height=280, showlegend=False,
            paper_bgcolor="#1e2535", plot_bgcolor="#1e2535",
            font_color="#aaa", margin=dict(t=20, b=20),
        )
        st.plotly_chart(fig3, use_container_width=True)

    st.markdown("#### Ultimas Cargas")
    df_cargas = get_cargas_historico(_eid())
    if len(df_cargas) > 0:
        st.dataframe(df_cargas.head(6), hide_index=True, use_container_width=True)


# ─────────────────────────────────────────────────────────────────────────────
# PAGINA: CARGAR CARTERA
# ─────────────────────────────────────────────────────────────────────────────

def page_cargar():
    st.markdown("## Cargar Cartera")
    st.markdown("*Sube el Excel del dia. El sistema detecta automaticamente clientes conocidos vs nuevos.*")
    st.divider()

    pipeline = load_model()
    if pipeline is None:
        st.error("Modelo no encontrado. Ejecuta primero: python src/main.py")
        return

    uploaded = st.file_uploader(
        "Selecciona el archivo de cartera — xlsx, xls, csv, txt, docx",
        type=None,
        help="Acepta cualquier formato: Excel (.xlsx/.xls), CSV, TXT, Google Docs (.docx)",
    )

    if uploaded is None:
        with st.expander("Columnas esperadas en el archivo"):
            col_l, col_r = st.columns(2)
            with col_l:
                st.markdown("""
| Columna | Descripcion |
|---|---|
| `cliente_id` | Identificador unico |
| `dpd` | Dias en mora |
| `saldo_total` | Saldo vencido |
| `bucket_mora` | B1, B2, B3, B4 |
| `rpc_rate` | Tasa de contacto (0.0-1.0) |
| `promesas_cumplidas` | Promesas honradas |
                """)
            with col_r:
                st.markdown("""
| Columna | Descripcion |
|---|---|
| `promesas_rotas` | Promesas incumplidas |
| `dias_ultimo_contacto` | Dias desde ultimo RPC |
| `ultimo_estado_marcado` | RPC_PROMESA, NO_CONTESTA... |
| `estado_laboral` | Dependiente, Independiente... |
| `ingreso_mensual` | Ingreso declarado |
| `ratio_deuda_ingreso` | Deuda / (ingreso x 12) |
                """)
        return

    try:
        df_raw = _read_any_file(uploaded)
    except Exception as e:
        st.error(f"Error al leer el archivo: {e}")
        return

    st.success(f"Archivo cargado: **{uploaded.name}** — {len(df_raw):,} registros | {df_raw.shape[1]} columnas")

    with st.expander("Vista previa (5 primeros registros)"):
        st.dataframe(df_raw.head(5), use_container_width=True)

    st.markdown("<br>", unsafe_allow_html=True)
    if st.button("PROCESAR CARTERA Y CALCULAR SCORES", type="primary", use_container_width=True):
        usuario = st.session_state["user"]["username"]
        prog = st.progress(0, "Verificando en base de datos...")

        try:
            prog.progress(25, "Identificando clientes conocidos vs nuevos...")
            df_result, stats = process_upload(df_raw, pipeline, usuario, uploaded.name)
            prog.progress(90, "Guardando en base de datos...")
            prog.progress(100, "Completado.")
        except Exception as e:
            st.error(f"Error durante el procesamiento: {e}")
            import traceback
            st.code(traceback.format_exc())
            return

        st.markdown("<br>", unsafe_allow_html=True)
        st.markdown("### Resultado del Procesamiento")

        # ── Fila 1: resumen del archivo ──────────────────────────────────────────
        r1, r2, r3, r4 = st.columns(4)
        with r1: card_metrica("Registros en archivo",    f"{stats['original']:,}",  color="#3b82f6")
        with r2: card_metrica("Duplicados eliminados",   f"{stats['duplicados']:,}", f"misma clienta, se quedo la mas reciente", "#f39c12")
        with r3: card_metrica("Cartera unica procesada", f"{stats['total']:,}",      color="#3b82f6")
        with r4: card_metrica("Cartera NUEVA (no estaba en BD)", f"{stats['nuevos']:,}",
                               f"{stats['nuevos']/max(stats['total'],1)*100:.1f}% de la carga", "#8b5cf6")

        st.markdown("<br>", unsafe_allow_html=True)

        # ── Fila 2: segmentos ────────────────────────────────────────────────────
        sc = df_result["segmento"].value_counts()
        s1, s2, s3 = st.columns(3)
        with s1: card_metrica("Segmento ALTO",  f"{sc.get('ALTO',0):,}",
                               "Llamada IVR / WhatsApp / SMS", "#27ae60")
        with s2: card_metrica("Segmento MEDIO", f"{sc.get('MEDIO',0):,}",
                               "Marcador predictivo + Agente", "#f39c12")
        with s3: card_metrica("Segmento BAJO",  f"{sc.get('BAJO',0):,}",
                               "Agente Senior + SMS urgente", "#e74c3c")

        st.divider()
        st.markdown("### Cartera con Plan Personalizado")
        st.caption("Cada clienta tiene su accion especifica segun score, mora, contactabilidad y promesas.")

        cols_show = [c for c in [
            "cliente_id", "score_operativo", "segmento", "plan_personalizado",
            "dpd", "saldo_total", "prob_pago", "bucket_mora", "estado_carga",
        ] if c in df_result.columns]

        st.dataframe(
            df_result[cols_show].sort_values("score_operativo", ascending=False).head(500),
            hide_index=True,
            use_container_width=True,
            column_config={
                "score_operativo":   st.column_config.ProgressColumn("Score", min_value=1, max_value=100),
                "prob_pago":         st.column_config.NumberColumn("Prob. Pago", format="%.1%"),
                "saldo_total":       st.column_config.NumberColumn("Saldo", format="S/ %.0f"),
                "plan_personalizado":st.column_config.TextColumn("Plan de Accion", width="large"),
            },
        )
        st.caption(f"Mostrando primeros 500 de {len(df_result):,} registros.")

        st.divider()
        ts = datetime.now().strftime("%Y%m%d_%H%M")
        d1, d2 = st.columns(2)
        with d1:
            st.download_button(
                "Descargar Excel Completo (4 hojas)",
                data=export_excel(df_result),
                file_name=f"cartera_{ts}.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                type="primary", use_container_width=True,
            )
        with d2:
            buf = io.StringIO()
            df_result.to_csv(buf, index=False)
            st.download_button(
                "Descargar CSV para Dialer",
                data=buf.getvalue(),
                file_name=f"dialer_{ts}.csv",
                mime="text/csv",
                use_container_width=True,
            )


# ─────────────────────────────────────────────────────────────────────────────
# PAGINA: ANALISIS
# ─────────────────────────────────────────────────────────────────────────────

def page_analisis():
    st.markdown("## Analisis de Cartera")
    st.divider()

    df = get_all_clientes_df(limit=20000, empresa_id=_eid())
    if len(df) == 0:
        st.info("Carga tu primera cartera para ver el analisis.")
        return

    segs = st.multiselect(
        "Filtrar por segmento", ["ALTO", "MEDIO", "BAJO"],
        default=["ALTO", "MEDIO", "BAJO"]
    )
    df = df[df["segmento"].isin(segs)]
    st.caption(f"{len(df):,} clientes en la seleccion actual")

    tab1, tab2, tab3 = st.tabs(["Score y Mora", "Contactabilidad", "Saldo y Riesgo"])

    chart_layout = dict(
        paper_bgcolor="#1e2535", plot_bgcolor="#1e2535",
        font_color="#aaa", margin=dict(t=30, b=20),
    )

    with tab1:
        c1, c2 = st.columns(2)
        with c1:
            st.markdown("##### Score vs Dias en Mora (DPD)")
            if "dpd" in df.columns:
                fig = px.scatter(
                    df.sample(min(2000, len(df))),
                    x="dpd", y="score_operativo",
                    color="segmento", color_discrete_map=COLORES,
                    opacity=0.55, template="plotly_dark",
                    labels={"dpd": "DPD", "score_operativo": "Score"},
                )
                fig.update_layout(**chart_layout, height=360)
                st.plotly_chart(fig, use_container_width=True)
        with c2:
            st.markdown("##### Clientes por Bucket de Mora")
            if "bucket_mora" in df.columns:
                fig2 = px.histogram(
                    df, x="bucket_mora", color="segmento",
                    color_discrete_map=COLORES, barmode="group",
                    template="plotly_dark",
                    category_orders={"bucket_mora": ["B1","B2","B3","B4"]},
                )
                fig2.update_layout(**chart_layout, height=360)
                st.plotly_chart(fig2, use_container_width=True)

    with tab2:
        c1, c2 = st.columns(2)
        with c1:
            st.markdown("##### RPC Rate por Segmento")
            if "rpc_rate" in df.columns:
                fig3 = px.box(
                    df, x="segmento", y="rpc_rate",
                    color="segmento", color_discrete_map=COLORES,
                    template="plotly_dark",
                    labels={"rpc_rate": "Tasa de Contacto Efectivo"},
                )
                fig3.update_layout(**chart_layout, height=360, showlegend=False)
                st.plotly_chart(fig3, use_container_width=True)
        with c2:
            st.markdown("##### Ultimo Estado de Marcado")
            if "ultimo_estado_marcado" in df.columns:
                cnt = df["ultimo_estado_marcado"].value_counts().head(7).reset_index()
                cnt.columns = ["Estado", "Clientes"]
                fig4 = px.bar(
                    cnt, x="Clientes", y="Estado",
                    orientation="h", template="plotly_dark",
                    color="Clientes", color_continuous_scale="Blues",
                )
                fig4.update_layout(
                    **chart_layout, height=360, showlegend=False,
                    yaxis={"categoryorder": "total ascending"},
                )
                st.plotly_chart(fig4, use_container_width=True)

    with tab3:
        c1, c2 = st.columns(2)
        with c1:
            st.markdown("##### Saldo Total por Segmento (S/)")
            if "saldo_total" in df.columns:
                s = df.groupby("segmento")["saldo_total"].sum().reset_index()
                fig5 = px.bar(
                    s, x="segmento", y="saldo_total",
                    color="segmento", color_discrete_map=COLORES,
                    text_auto=".3s", template="plotly_dark",
                )
                fig5.update_layout(**chart_layout, height=360, showlegend=False)
                st.plotly_chart(fig5, use_container_width=True)
        with c2:
            st.markdown("##### Probabilidad de Pago por Segmento")
            fig6 = px.box(
                df, x="segmento", y="prob_pago",
                color="segmento", color_discrete_map=COLORES,
                template="plotly_dark",
                labels={"prob_pago": "Probabilidad de Pago"},
            )
            fig6.update_layout(**chart_layout, height=360, showlegend=False)
            st.plotly_chart(fig6, use_container_width=True)


# ─────────────────────────────────────────────────────────────────────────────
# PAGINA: HISTORIAL
# ─────────────────────────────────────────────────────────────────────────────

def page_historial():
    st.markdown("## Historial de Cargas")
    st.divider()

    df = get_cargas_historico(_eid())
    if len(df) == 0:
        st.info("No hay cargas registradas aun.")
        return

    c1, c2, c3 = st.columns(3)
    with c1: card_metrica("Total de Cargas", f"{len(df):,}", color="#3b82f6")
    with c2: card_metrica("Total Procesados", f"{df['total_registros'].sum():,}", color="#27ae60")
    with c3: card_metrica("Clientes Nuevos Registrados", f"{df['registros_nuevos'].sum():,}", color="#8b5cf6")

    st.markdown("<br>", unsafe_allow_html=True)
    st.dataframe(
        df, hide_index=True, use_container_width=True,
        column_config={
            "total_registros":        st.column_config.NumberColumn("Total"),
            "registros_nuevos":       st.column_config.NumberColumn("Nuevos"),
            "registros_actualizados": st.column_config.NumberColumn("Actualizados"),
        },
    )


# ─────────────────────────────────────────────────────────────────────────────
# PAGINA: ESTRATEGIAS
# ─────────────────────────────────────────────────────────────────────────────

def page_estrategias():
    st.markdown("## Estrategias de Cobranza de Clase Mundial")
    st.markdown("*Basadas en metodologias de Hoist Finance, Encore Capital, FICO TRIAD, COFACE e Intrum.*")
    st.divider()

    score_rng = {"ALTO": "67 - 100", "MEDIO": "34 - 66", "BAJO": "1 - 33"}
    nombres   = {"ALTO": "ALTO", "MEDIO": "MEDIO", "BAJO": "BAJO"}

    for seg, color in [("ALTO", "#27ae60"), ("MEDIO", "#f39c12"), ("BAJO", "#e74c3c")]:
        e = ESTRATEGIAS[seg]
        st.markdown(f"""
        <div class="card-estrategia" style="border-left: 4px solid {color}">
            <div style="display:flex; align-items:center; gap:12px; margin-bottom:14px">
                <span style="font-size:1.2rem; font-weight:800; color:{color}">
                    Segmento {nombres[seg]}
                </span>
                <span style="font-size:0.8rem; color:#6b7a99; background:#0f1117;
                    padding:3px 10px; border-radius:20px; border:1px solid #2a3045">
                    Score {score_rng[seg]}
                </span>
            </div>
            <div style="display:grid; grid-template-columns:1fr 1fr; gap:12px">
                <div>
                    <div style="color:#6b7a99; font-size:0.75rem; margin-bottom:2px">CANAL</div>
                    <div style="color:#e2e8f0; font-size:0.9rem">{e['canal']}</div>
                </div>
                <div>
                    <div style="color:#6b7a99; font-size:0.75rem; margin-bottom:2px">ACCION</div>
                    <div style="color:#e2e8f0; font-size:0.9rem">{e['accion']}</div>
                </div>
                <div>
                    <div style="color:#6b7a99; font-size:0.75rem; margin-bottom:2px">OFERTA</div>
                    <div style="color:#e2e8f0; font-size:0.9rem">{e['oferta']}</div>
                </div>
                <div>
                    <div style="color:#6b7a99; font-size:0.75rem; margin-bottom:2px">FRECUENCIA</div>
                    <div style="color:#e2e8f0; font-size:0.9rem">{e['frecuencia']}</div>
                </div>
                <div>
                    <div style="color:#6b7a99; font-size:0.75rem; margin-bottom:2px">KPIs OBJETIVO</div>
                    <div style="color:{color}; font-size:0.85rem; font-weight:600">{e['kpis']}</div>
                </div>
                <div>
                    <div style="color:#6b7a99; font-size:0.75rem; margin-bottom:2px">ESCALACION</div>
                    <div style="color:#e2e8f0; font-size:0.9rem">{e['escalacion']}</div>
                </div>
            </div>
        </div>
        """, unsafe_allow_html=True)

        with st.expander(f"Script / Guion para agentes — Segmento {seg}"):
            st.info(e["script"])
            st.caption(f"Referencia: {e['referencia']}")

    st.divider()
    st.markdown("#### Comparacion Rapida")
    comp = pd.DataFrame([
        {"Metrica": "Canal",               "ALTO": "WhatsApp / SMS / IVR",     "MEDIO": "Marcador Predictivo + Agente", "BAJO": "Agente Senior + SMS/WhatsApp"},
        {"Metrica": "Costo por contacto",  "ALTO": "S/ 0.20-0.50",             "MEDIO": "S/ 2.00-3.50",                "BAJO": "S/ 5.00-10.00"},
        {"Metrica": "Frecuencia maxima",   "ALTO": "1 WA + 1 SMS / semana",    "MEDIO": "3 llamadas / dia",            "BAJO": "1 llamada diaria + SMS"},
        {"Metrica": "Conversion objetivo", "ALTO": "Respuesta WA >= 25%",      "MEDIO": "PTP >= 30%",                  "BAJO": "Settlement >= 15%"},
        {"Metrica": "Mejor horario",       "ALTO": "8am-10pm (WA/SMS)",        "MEDIO": "9-11am y 6-8pm L-V",         "BAJO": "9am-6pm L-V"},
        {"Metrica": "ROI estimado",        "ALTO": "400-600%",                  "MEDIO": "150-300%",                    "BAJO": "50-150%"},
    ])
    st.dataframe(comp, hide_index=True, use_container_width=True)


# ─────────────────────────────────────────────────────────────────────────────
# PAGINA: ADMIN
# ─────────────────────────────────────────────────────────────────────────────

def page_admin():
    if st.session_state["user"]["rol"] != "admin":
        st.error("Acceso denegado. Solo para Administradores.")
        return

    st.markdown("## Panel de Administracion")
    st.divider()

    _tabs = ["Usuarios del Sistema", "Crear Usuario", "Cambiar Contrasena"]
    if _eid() == 1:
        _tabs.append("Gestionar Empresas")
    _tab_objs = st.tabs(_tabs)
    tab1, tab2, tab3 = _tab_objs[0], _tab_objs[1], _tab_objs[2]

    with tab1:
        st.markdown("#### Lista de Usuarios")
        df_u = get_all_users(_eid())
        st.dataframe(
            df_u, hide_index=True, use_container_width=True,
            column_config={
                "activo": st.column_config.CheckboxColumn("Activo"),
                "rol":    st.column_config.SelectboxColumn("Rol", options=["admin","colaborador"]),
            },
        )
        st.markdown("#### Activar / Desactivar Usuario")
        cur = st.session_state["user"]["username"]
        opciones = [u for u in df_u["username"].tolist() if u != cur]
        if opciones:
            sel = st.selectbox("Selecciona usuario", opciones)
            if st.button(f"Cambiar estado de '{sel}'", type="primary"):
                toggle_user_status(sel, _eid())
                st.success(f"Estado de '{sel}' actualizado.")
                st.rerun()

    with tab2:
        st.markdown("#### Nuevo Usuario")
        with st.form("crear_usuario"):
            c1, c2 = st.columns(2)
            nu = c1.text_input("Username")
            nn = c2.text_input("Nombre completo")
            ne = c1.text_input("Email")
            nr = c2.selectbox("Rol", ["colaborador", "admin"])
            np_ = st.text_input("Contrasena", type="password")
            nc  = st.text_input("Confirmar contrasena", type="password")
            if st.form_submit_button("Crear Usuario", type="primary"):
                if np_ != nc:
                    st.error("Las contrasenas no coinciden.")
                elif not nu or not nn:
                    st.error("Username y nombre son obligatorios.")
                else:
                    ok, msg = create_user(nu, np_, nn, ne, nr, _eid())
                    (st.success if ok else st.error)(msg)

    with tab3:
        st.markdown("#### Cambiar Contrasena")
        df_u2 = get_all_users(_eid())
        with st.form("cambiar_pass"):
            us = st.selectbox("Usuario", df_u2["username"].tolist() if len(df_u2) > 0 else [])
            p1 = st.text_input("Nueva contrasena", type="password")
            p2 = st.text_input("Confirmar contrasena", type="password")
            if st.form_submit_button("Actualizar Contrasena", type="primary"):
                if p1 != p2:
                    st.error("Las contrasenas no coinciden.")
                else:
                    ok, msg = update_password(us, p1, _eid())
                    (st.success if ok else st.error)(msg)

    if _eid() == 1 and len(_tab_objs) > 3:
        with _tab_objs[3]:
            st.markdown("#### Empresas registradas")
            df_emp = get_all_empresas()
            st.dataframe(df_emp, hide_index=True, use_container_width=True)
            st.markdown("#### Agregar nueva empresa")
            with st.form("nueva_empresa"):
                e_nombre = st.text_input("Nombre de la empresa")
                e_slug   = st.text_input("Slug (identificador único, sin espacios)", placeholder="empresa-abc")
                if st.form_submit_button("Crear Empresa", type="primary"):
                    if not e_nombre or not e_slug:
                        st.error("Completa nombre y slug.")
                    else:
                        ok, msg = create_empresa(e_nombre, e_slug.lower().replace(" ", "-"))
                        (st.success if ok else st.error)(msg)
                        if ok:
                            st.info(f"Empresa '{e_nombre}' creada. Ahora crea un usuario admin para esa empresa desde esta misma página seleccionando el empresa_id correspondiente.")


# ─────────────────────────────────────────────────────────────────────────────
# PAGINA: REPORTE VICIDIAL
# ─────────────────────────────────────────────────────────────────────────────

def _read_any_file(uploaded) -> pd.DataFrame:
    """Lee CSV, Excel, TXT o DOCX y devuelve un DataFrame."""
    name = uploaded.name.lower()

    if name.endswith((".xlsx", ".xls")):
        return pd.read_excel(uploaded)

    if name.endswith(".docx"):
        from docx import Document
        doc = Document(uploaded)
        # Buscar tablas dentro del documento
        if doc.tables:
            tbl = doc.tables[0]
            headers = [cell.text.strip() for cell in tbl.rows[0].cells]
            data = []
            for row in tbl.rows[1:]:
                data.append([cell.text.strip() for cell in row.cells])
            df = pd.DataFrame(data, columns=headers)
            # Limpiar columnas vacías
            df = df.loc[:, df.columns.str.strip() != ""]
            df = df[df.apply(lambda r: r.str.strip().ne("").any(), axis=1)]
            return df
        # Si no hay tablas, intentar extraer texto como CSV/TSV
        lines = [p.text for p in doc.paragraphs if p.text.strip()]
        if lines:
            from io import StringIO
            text = "\n".join(lines)
            for sep in ["\t", ";", ",", "|"]:
                try:
                    df = pd.read_csv(StringIO(text), sep=sep)
                    if len(df.columns) > 1:
                        return df
                except Exception:
                    pass
        raise ValueError("No se encontraron tablas ni datos estructurados en el documento.")

    # CSV y TXT: detectar separador y encoding automáticamente
    for enc in ["utf-8", "latin-1", "cp1252"]:
        for sep in [",", ";", "\t", "|"]:
            try:
                df = pd.read_csv(uploaded, encoding=enc, sep=sep, low_memory=False)
                if len(df.columns) > 1:
                    return df
                uploaded.seek(0)
            except Exception:
                uploaded.seek(0)
    return pd.read_csv(uploaded, encoding="latin-1", low_memory=False)


def _vic_load(uploaded) -> pd.DataFrame:
    return _vic_load_with_legend(uploaded)[0]


def _vic_load_with_legend(uploaded) -> tuple:
    """Lee el archivo de llamadas y, si el Excel trae una hoja extra con
    leyenda de códigos (columnas tipo Código/Descripción), la devuelve como dict."""
    name = uploaded.name.lower()
    if name.endswith((".xlsx", ".xls")):
        xl = pd.ExcelFile(uploaded)
        if len(xl.sheet_names) > 1:
            best, best_cols = None, -1
            for s in xl.sheet_names:
                sdf = xl.parse(s, nrows=3)
                if _vic_find(sdf, "status") and sdf.shape[1] > best_cols:
                    best, best_cols = s, sdf.shape[1]
            if best is None:
                best = max(xl.sheet_names, key=lambda s: xl.parse(s, nrows=1).shape[1])
            legend = {}
            for s in xl.sheet_names:
                if s == best:
                    continue
                ldf = xl.parse(s)
                col_cod  = _vic_find(ldf, "codigo", "código", "status", "code")
                col_desc = _vic_find(ldf, "descripcion", "descripción", "desc")
                if col_cod and col_desc:
                    for _, row in ldf.iterrows():
                        cod = str(row[col_cod]).strip().upper()
                        desc = str(row[col_desc]).strip()
                        if cod and cod != "NAN" and desc and desc.upper() != "NAN":
                            legend[cod] = desc
            return xl.parse(best), legend
        return xl.parse(xl.sheet_names[0]), {}
    return _read_any_file(uploaded), {}


def _vic_fmt(seconds) -> str:
    if seconds is None or (isinstance(seconds, float) and np.isnan(seconds)):
        return "–"
    m, s = divmod(int(seconds), 60)
    return f"{m}m {s:02d}s"


def _vic_pct(num, den) -> str:
    return "0%" if den == 0 else f"{num/den*100:.1f}%"


def _vic_kpi(col, label, value, delta="", color="#3b82f6"):
    with col:
        st.markdown(f"""
        <div style="background:#1e2535;border-radius:10px;padding:16px 18px;
                    border-left:4px solid {color};margin-bottom:8px">
          <div style="color:#8899aa;font-size:12px;text-transform:uppercase">{label}</div>
          <div style="color:#fff;font-size:26px;font-weight:700">{value}</div>
          <div style="color:#6b7a99;font-size:12px;margin-top:2px">{delta}</div>
        </div>""", unsafe_allow_html=True)


def _vic_alert(title, body, color="#dc2626"):
    st.markdown(f"""
    <div style="background:#1c1917;border:1px solid {color};border-radius:10px;
                padding:14px 18px;margin-bottom:8px">
      <div style="color:{color};font-weight:700;font-size:14px">⚠ {title}</div>
      <div style="color:#d1d5db;font-size:13px;margin-top:4px">{body}</div>
    </div>""", unsafe_allow_html=True)


def _vic_critical(total, contacted, machines, no_answer, avg_dur, agent_df):
    points = []
    cr = contacted / total if total else 0
    if cr < 0.20:
        points.append(("rojo", "Tasa de contacto CRÍTICA",
            f"Solo {cr*100:.1f}% de llamadas contactan (meta mínima 20%). "
            "Revisar base, horarios y estrategia de marcado."))
    elif cr < 0.35:
        points.append(("naranja", "Tasa de contacto por debajo del objetivo",
            f"{cr*100:.1f}% de contacto. Optimizar franjas horarias y depurar números inválidos."))

    mr = machines / total if total else 0
    if mr > 0.30:
        points.append(("naranja", "Alto porcentaje de contestadores automáticos",
            f"{mr*100:.1f}% caen en contestador/AM. Activar AMD o ajustar horarios."))

    if avg_dur is not None:
        if avg_dur < 30:
            points.append(("rojo", "Duración promedio muy corta",
                f"Promedio {avg_dur:.0f}s indica llamadas incompletas o desconexiones. "
                "Verificar conectividad y scripts de apertura."))
        elif avg_dur > 600:
            points.append(("naranja", "Llamadas con duración excesiva",
                f"Promedio {avg_dur/60:.1f} min. Revisar manejo de objeciones y scripts."))

    if agent_df is not None and "tasa_contacto" in agent_df.columns and len(agent_df) > 1:
        low = agent_df[agent_df["tasa_contacto"] < 15]
        if len(low) > 0:
            names = ", ".join(low.index.astype(str)[:5])
            points.append(("naranja", "Agentes con tasa de contacto < 15%",
                f"{names}. Requieren coaching o revisión de su base asignada."))

    if not points:
        points.append(("verde", "Sin alertas críticas detectadas",
            "El reporte no muestra problemas graves. Monitorear tendencias para mejora continua."))
    return points


def _vic_find(df: pd.DataFrame, *substrings) -> str | None:
    """Busca una columna cuyo nombre (en minúsculas) contenga alguno de los substrings."""
    cl = {c.lower().strip(): c for c in df.columns}
    for sub in substrings:
        for k, v in cl.items():
            if sub in k:
                return v
    return None


def _vic_status_norm(series: pd.Series) -> pd.Series:
    s = series.astype(str).str.strip().str.upper()
    s = s.str.replace(r"\.0$", "", regex=True)
    return s.where(~s.str.fullmatch(r"\d"), s.str.zfill(2))


def _vic_extract_date_from_name(name: str):
    import re
    m = re.search(r"(20\d{6})", name)
    if m:
        try:
            return datetime.strptime(m.group(1), "%Y%m%d").date()
        except Exception:
            return None
    return None


def _vic_export_cols(df: pd.DataFrame, monto_subs=("postal_code",), estado_subs=("first_name",)) -> dict:
    return {
        "status":       _vic_find(df, "status"),
        "phone":        _vic_find(df, "phone_number", "phone", "telefono"),
        "date":         _vic_find(df, "call_date", "date", "fecha"),
        "agent":        _vic_find(df, "user", "agent", "agente"),
        "entidad":      _vic_find(df, "list_id", "campaign_id", "entidad"),
        "monto":        _vic_find(df, *monto_subs),
        "estado_deudor":_vic_find(df, *estado_subs),
        "lead_id":      _vic_find(df, "lead_id"),
    }


def _vic_filter_jornada(df: pd.DataFrame, c: dict, fecha):
    """Si la fecha es sábado, filtra la jornada a 8:00-12:00 (media jornada)."""
    if fecha is None or fecha.weekday() != 5 or not c["date"]:
        return df, False
    dt = pd.to_datetime(df[c["date"]], errors="coerce")
    mask = dt.dt.hour.between(8, 11)
    return df[mask], True


def _vic_tablero_contactabilidad(df: pd.DataFrame, c: dict, fecha, promesa=None, saludo_st=None,
                                  humanos_st=None, labels=None) -> dict:
    promesa    = promesa if promesa is not None else _VIC_GRAL_PROMESA
    saludo_st  = saludo_st if saludo_st is not None else _VIC_GRAL_SALUDO
    humanos_st = humanos_st if humanos_st is not None else _VIC_GRAL_HUMANOS
    labels     = labels if labels is not None else _VIC_GRAL_LABELS
    total = len(df)
    humanos  = df[df["_st"].isin(humanos_st)]
    promesas = df[df["_st"].isin(promesa)]
    saludo   = df[df["_st"] == saludo_st]
    gestion  = humanos[humanos["_st"] != saludo_st]

    monto_total = 0.0
    if c["monto"]:
        monto_total = pd.to_numeric(promesas[c["monto"]], errors="coerce").fillna(0).sum()

    promesa_lbl = "+".join(sorted(promesa))
    resumen = pd.DataFrame({
        "Métrica": ["Fecha", "Total llamadas", "Total humanos (contacto)", "% Contactabilidad",
                    f"Cuelga en saludo (status {saludo_st})", "% Cuelga en saludo",
                    "Gestión efectiva (humanos - saludo)",
                    f"Promesas de pago ({promesa_lbl})", "% Promesas sobre gestión",
                    "Monto comprometido (S/)"],
        "Valor": [str(fecha), total, len(humanos), _vic_pct(len(humanos), total),
                  len(saludo), _vic_pct(len(saludo), total),
                  len(gestion),
                  len(promesas), _vic_pct(len(promesas), len(gestion)),
                  round(float(monto_total), 2)],
    })

    por_estado = df["_st"].value_counts().rename_axis("Status").reset_index(name="Llamadas")
    por_estado["Descripción"] = por_estado["Status"].map(labels).fillna(
        por_estado["Status"].map(_VIC_STATUS_LABELS)).fillna("—")
    por_estado["%"] = (por_estado["Llamadas"] / total * 100).round(2) if total else 0

    por_entidad = pd.DataFrame()
    if c["entidad"]:
        ent = df.copy()
        ent["_evasion"]  = ent["_st"].isin(_VIC_NOANSWER | _VIC_MACHINE)
        ent["_saludo"]   = ent["_st"] == saludo_st
        ent["_humano"]   = ent["_st"].isin(humanos_st)
        ent["_promesa"]  = ent["_st"].isin(promesa)
        por_entidad = ent.groupby(c["entidad"]).agg(
            total=("_st", "count"),
            humanos=("_humano", "sum"),
            evasion=("_evasion", "sum"),
            cuelga_saludo=("_saludo", "sum"),
            promesas=("_promesa", "sum"),
        ).reset_index()
        por_entidad["%_evasion"] = (por_entidad["evasion"] / por_entidad["total"] * 100).round(1)
        por_entidad["%_cuelga"]  = (por_entidad["cuelga_saludo"] / por_entidad["total"] * 100).round(1)

    return {
        "resumen": resumen, "por_estado": por_estado, "por_entidad": por_entidad,
        "kpis": {"total": total, "humanos": len(humanos), "promesas": len(promesas),
                 "saludo": len(saludo), "monto": float(monto_total), "gestion": len(gestion)},
    }


def _vic_control_recontacto(df: pd.DataFrame, c: dict, fecha, promesa=None, labels=None) -> dict | None:
    promesa = promesa if promesa is not None else _VIC_GRAL_PROMESA
    labels  = labels if labels is not None else _VIC_GRAL_LABELS
    key = c["phone"] or c["lead_id"]
    if not key:
        return None
    g = df.groupby(key)
    res = g.size().rename("llamadas").reset_index()
    res["ultimo_estado"] = g["_st"].last().values
    if c["monto"]:
        res["monto"] = pd.to_numeric(g[c["monto"]].last(), errors="coerce").fillna(0).values
    if c["estado_deudor"]:
        res["estado_deudor"] = g[c["estado_deudor"]].last().values
    res["requiere_recontacto"] = ~res["ultimo_estado"].isin(promesa)
    res["descripcion_estado"]  = res["ultimo_estado"].map(labels).fillna(
        res["ultimo_estado"].map(_VIC_STATUS_LABELS)).fillna("—")

    con_promesa = int((~res["requiere_recontacto"]).sum())
    pendientes  = int(res["requiere_recontacto"].sum())
    resumen = pd.DataFrame({
        "Métrica": ["Fecha", "Total leads contactados", "Con promesa de pago",
                    "Pendientes de recontacto", "% Pendientes de recontacto"],
        "Valor": [str(fecha), len(res), con_promesa, pendientes, _vic_pct(pendientes, len(res))],
    })
    return {"resumen": resumen, "detalle": res,
            "kpis": {"total_leads": len(res), "promesas": con_promesa, "pendientes": pendientes}}


def _vic_tipificacion_gestion(df: pd.DataFrame, c: dict, fecha, labels=None) -> dict:
    labels = labels if labels is not None else _VIC_GRAL_LABELS
    total = len(df)
    g = df.groupby("_st").size().rename("llamadas").reset_index().rename(columns={"_st": "Status"})
    g["%"] = (g["llamadas"] / total * 100).round(2) if total else 0
    g["Descripción"] = g["Status"].map(labels).fillna(g["Status"].map(_VIC_STATUS_LABELS)).fillna("—")
    if c["monto"]:
        montos = df.groupby("_st")[c["monto"]].apply(lambda s: pd.to_numeric(s, errors="coerce").fillna(0).sum())
        g["monto_total"] = g["Status"].map(montos).fillna(0).round(2)

    por_estado_deudor = pd.DataFrame()
    if c["estado_deudor"]:
        por_estado_deudor = df.groupby(c["estado_deudor"]).size().rename("count") \
            .reset_index().sort_values("count", ascending=False)

    resumen = pd.DataFrame({
        "Métrica": ["Fecha", "Total registros", "Tipos de gestión distintos"],
        "Valor": [str(fecha), total, df["_st"].nunique()],
    })
    return {"resumen": resumen, "por_status": g, "por_estado_deudor": por_estado_deudor,
            "kpis": {"total": total}}


def _vic_append_historico(prev_upload, resumen_df: pd.DataFrame) -> pd.DataFrame:
    row_df = pd.DataFrame([dict(zip(resumen_df["Métrica"], resumen_df["Valor"]))])
    hist = row_df
    if prev_upload is not None:
        try:
            prev_upload.seek(0)
            hist_prev = pd.read_excel(prev_upload, sheet_name="Histórico")
            hist = pd.concat([hist_prev, row_df], ignore_index=True)
        except Exception:
            try:
                prev_upload.seek(0)
                hist_prev = pd.read_excel(prev_upload, sheet_name=0)
                hist = pd.concat([hist_prev, row_df], ignore_index=True)
            except Exception:
                hist = row_df
    if "Fecha" in hist.columns:
        hist = hist.drop_duplicates(subset=["Fecha"], keep="last").reset_index(drop=True)
    return hist


def _vic_write_excel(sheets: dict) -> io.BytesIO:
    buf = io.BytesIO()
    with pd.ExcelWriter(buf, engine="openpyxl") as writer:
        for name, sdf in sheets.items():
            if sdf is not None and len(sdf) > 0:
                sdf.to_excel(writer, sheet_name=name[:31], index=False)
    buf.seek(0)
    return buf


def _vic_trend_line(hist: pd.DataFrame, metric: str, label: str, fmt="{:.1f}") -> str | None:
    if hist is None or metric not in hist.columns or len(hist) < 2:
        return None
    try:
        prev = float(str(hist.iloc[-2][metric]).replace("%", ""))
        curr = float(str(hist.iloc[-1][metric]).replace("%", ""))
    except (ValueError, TypeError):
        return None
    delta = curr - prev
    arrow = "▲" if delta > 0 else ("▼" if delta < 0 else "→")
    return f"{label}: {fmt.format(curr)} ({arrow} {fmt.format(abs(delta))} vs. día anterior)"


def _vic_resumen_color(metrica: str) -> str:
    m = str(metrica).lower()
    if "promesa" in m:
        return "#22c55e"
    if "monto" in m:
        return "#a855f7"
    if "cuelga" in m or "saludo" in m:
        return "#ef4444"
    if "contactabilidad" in m or "gestión efectiva" in m or "humanos" in m or "pendientes" in m:
        return "#3b82f6"
    return "#64748b"


def _vic_estado_cards(df: pd.DataFrame, promesa, saludo_st, cols_per_row: int = 4, top: int = 16):
    """Muestra cada status como tarjeta (código, descripción, llamadas, %), coloreada
    según si es promesa de pago (verde), problema operativo (rojo/naranja) u otro (azul)."""
    rows = df.sort_values("Llamadas", ascending=False).head(top).to_dict("records")
    for i in range(0, len(rows), cols_per_row):
        chunk = rows[i:i + cols_per_row]
        cols = st.columns(len(chunk))
        for col, r in zip(cols, chunk):
            color = _coq_status_color(r.get("Status", ""), promesa, saludo_st)
            desc = r.get("Descripción", "—")
            with col:
                st.markdown(f"""
                <div style="background:#1a2333;border-radius:10px;padding:14px 16px;
                            border-left:4px solid {color};margin-bottom:10px;min-height:118px">
                  <div style="display:flex;justify-content:space-between;align-items:baseline">
                    <div style="color:#fff;font-size:17px;font-weight:800">{r.get('Status', '')}</div>
                    <div style="color:{color};font-size:15px;font-weight:700">{r.get('%', 0)}%</div>
                  </div>
                  <div style="color:#9fb0cc;font-size:12.5px;margin-top:4px;line-height:1.3">{desc}</div>
                  <div style="color:#6b7a99;font-size:12px;margin-top:8px">{int(r.get('Llamadas', 0)):,} llamadas</div>
                </div>""", unsafe_allow_html=True)


def _vic_entidad_cards(df: pd.DataFrame, ent_col: str, cols_per_row: int = 4, top: int = 12):
    """Muestra cada entidad (lista/cartera) como una tarjeta con su % de evasión y % de cuelga,
    en vez de un gráfico de barras (más fácil de leer con códigos numéricos largos)."""
    rows = df.sort_values("total", ascending=False).head(top).to_dict("records")
    for i in range(0, len(rows), cols_per_row):
        chunk = rows[i:i + cols_per_row]
        cols = st.columns(len(chunk))
        for col, r in zip(cols, chunk):
            evasion = r.get("%_evasion", 0)
            cuelga  = r.get("%_cuelga", 0)
            color = "#ef4444" if evasion > 40 else "#f59e0b" if evasion > 20 else "#22c55e"
            with col:
                st.markdown(f"""
                <div style="background:#1a2333;border-radius:10px;padding:14px 16px;
                            border-left:4px solid {color};margin-bottom:10px">
                  <div style="color:#8899aa;font-size:11px;text-transform:uppercase;
                              letter-spacing:.03em">Entidad / Lista</div>
                  <div style="color:#fff;font-size:18px;font-weight:700;margin-top:2px">{r.get(ent_col, '')}</div>
                  <div style="color:#6b7a99;font-size:12px;margin-top:6px">{int(r.get('total', 0)):,} llamadas</div>
                  <div style="display:flex;gap:14px;margin-top:8px">
                    <div><span style="color:#ef4444;font-weight:700">{evasion:.1f}%</span>
                      <span style="color:#8899aa;font-size:11px"> evasión</span></div>
                    <div><span style="color:#f59e0b;font-weight:700">{cuelga:.1f}%</span>
                      <span style="color:#8899aa;font-size:11px"> cuelga</span></div>
                  </div>
                </div>""", unsafe_allow_html=True)


def _vic_resumen_cards(df: pd.DataFrame, cols_per_row: int = 5):
    """Renderiza la tabla Resumen como tarjetas con borde de color (mismo estilo que los KPIs),
    en vez de una tabla plana, para que sea más fácil de leer de un vistazo."""
    rows = df.to_dict("records")
    for i in range(0, len(rows), cols_per_row):
        chunk = rows[i:i + cols_per_row]
        cols = st.columns(len(chunk))
        for col, r in zip(cols, chunk):
            color = _vic_resumen_color(r.get("Métrica", ""))
            with col:
                st.markdown(f"""
                <div style="background:#1a2333;border-radius:10px;padding:14px 16px;
                            border-left:4px solid {color};margin-bottom:10px;min-height:84px">
                  <div style="color:#8899aa;font-size:11px;text-transform:uppercase;
                              letter-spacing:.03em">{r.get('Métrica', '')}</div>
                  <div style="color:#fff;font-size:21px;font-weight:700;margin-top:4px">{r.get('Valor', '')}</div>
                </div>""", unsafe_allow_html=True)


def page_vicidial():
    st.markdown("## 📞 Reportes Diarios — Campaña GRAL")
    st.markdown("*Sube los 6 archivos del día y genera los 3 reportes actualizados: "
                 "Tablero de Contactabilidad, Control de Recontacto y Tipificación de Gestión.*")
    st.divider()

    with st.expander("ℹ️ Reglas aplicadas en el cálculo"):
        st.markdown("""
- **Promesa de pago** = status `04` + `21` (siempre se suman ambos).
- **Status `01`** = cuelga en saludo (rechazo temprano), no cuenta como gestión.
- **Humanos** = status `01, 02, 04, 09, 14, 18, 19, 21` → hoja *Por Entidad* (evasión y cuelga por estado).
- **Sábado** = media jornada (8 AM–12 PM); se compara sábado contra sábado.
- El día nuevo se agrega al **acumulado (Histórico)** de cada reporte, leyendo el tablero del día anterior.
- **Monto comprometido** sale del campo `postal_code` y **estado del deudor** del campo `first_name`
  (campos reutilizados del export).
        """)

    st.markdown("#### 1. Archivos frescos de VICIdial (jornada del día) — 3 archivos")
    f_amd     = st.file_uploader("AST_AMD_log_report (.csv)", type=None, key="vic_amd")
    f_vdad    = st.file_uploader("AST_VDADstats (.csv)", type=None, key="vic_vdad")
    f_export  = st.file_uploader(
        "EXPORT_CALL_REPORT — Estados = ALL (.txt/.csv)  ⭐ insumo principal",
        type=None, key="vic_export")

    st.markdown("#### 2. Tableros del día anterior (acumulado) — 3 archivos")
    f_contact_prev    = st.file_uploader("Tablero_Contactabilidad_GRAL (.xlsx)", type=None, key="vic_contact_prev")
    f_recontacto_prev = st.file_uploader("Control_Recontacto_GRAL (.xlsx)", type=None, key="vic_recontacto_prev")
    f_tipif_prev      = st.file_uploader("Tipificacion_Gestion_GRAL (.xlsx)", type=None, key="vic_tipif_prev")

    with st.expander("➕ Extra (no cuenta en los 6) — diagnóstico de sobre-marcado"):
        f_carrier = st.file_uploader(
            "AST_carrier_log_report (.csv) — opcional, diagnóstico de sobre-marcado por troncal",
            type=None, key="vic_carrier")

    st.divider()

    if f_export is None:
        st.info("Sube al menos el **EXPORT_CALL_REPORT (Estados = ALL)** para generar los 3 reportes.")
        return

    if not st.button("🚀 Generar Reportes", type="primary"):
        return

    with st.spinner("Procesando..."):
        try:
            df = _vic_load(f_export)
        except Exception as e:
            st.error(f"No se pudo leer EXPORT_CALL_REPORT: {e}")
            return

        c = _vic_export_cols(df)
        if not c["status"]:
            st.error("No se encontró la columna **status** en EXPORT_CALL_REPORT. "
                     "Verifica que el export incluya esa columna.")
            return
        df["_st"] = _vic_status_norm(df[c["status"]])

        fecha = _vic_extract_date_from_name(f_export.name) or date.today()
        df, es_sabado = _vic_filter_jornada(df, c, fecha)
        if es_sabado:
            st.info("📅 Detectado **sábado** — se filtró la jornada a 8:00–12:00 (media jornada). "
                    "Compara este reporte contra el sábado anterior en la hoja *Histórico*.")

        contact = _vic_tablero_contactabilidad(df, c, fecha)
        recont  = _vic_control_recontacto(df, c, fecha)
        tipif   = _vic_tipificacion_gestion(df, c, fecha)

        hist_contact = _vic_append_historico(f_contact_prev, contact["resumen"])
        hist_tipif   = _vic_append_historico(f_tipif_prev, tipif["resumen"])
        hist_recont  = _vic_append_historico(f_recontacto_prev, recont["resumen"]) if recont else None

    # ── KPIs ──────────────────────────────────────────────────────────────────
    st.markdown("#### Resumen del Día")
    k1, k2, k3, k4, k5 = st.columns(5)
    kp = contact["kpis"]
    cr_color = "#27ae60" if kp["humanos"]/kp["total"] >= 0.35 else "#f39c12" if kp["humanos"]/kp["total"] >= 0.20 else "#e74c3c"
    _vic_kpi(k1, "Total llamadas", f"{kp['total']:,}")
    _vic_kpi(k2, "Contactabilidad (humanos)", _vic_pct(kp["humanos"], kp["total"]), color=cr_color)
    _vic_kpi(k3, "Promesas de pago (04+21)", f"{kp['promesas']:,}")
    _vic_kpi(k4, "Monto comprometido", f"S/ {kp['monto']:,.2f}")
    sal_color = "#e74c3c" if kp["saludo"]/kp["total"] > 0.40 else "#f39c12" if kp["saludo"]/kp["total"] > 0.25 else "#27ae60"
    _vic_kpi(k5, "Cuelga en saludo (01)", _vic_pct(kp["saludo"], kp["total"]), color=sal_color)

    # ── Alertas / tendencia ──────────────────────────────────────────────────
    st.markdown("#### Alertas y Tendencia")
    saludo_rate = kp["saludo"] / kp["total"] if kp["total"] else 0
    if saludo_rate > 0.40:
        _vic_alert("Evasión / cuelga en saludo CRÍTICA",
                   f"{saludo_rate*100:.1f}% de las llamadas cuelgan en el saludo (status 01). "
                   "Revisar guion de apertura, horario de marcado y calidad de la base.", "#dc2626")
    elif saludo_rate > 0.25:
        _vic_alert("Cuelga en saludo elevado",
                   f"{saludo_rate*100:.1f}% cuelgan en el saludo. Vigilar tendencia.", "#f59e0b")

    if f_vdad is not None:
        try:
            vdf = _vic_load(f_vdad)
            drop_col  = _vic_find(vdf, "drop")
            calls_col = _vic_find(vdf, "call")
            if drop_col and calls_col:
                drops = pd.to_numeric(vdf[drop_col], errors="coerce").fillna(0).sum()
                calls = pd.to_numeric(vdf[calls_col], errors="coerce").fillna(0).sum()
                if calls > 0:
                    drop_rate = drops / calls * 100
                    if drop_rate > 5:
                        _vic_alert("DROP / sobre-marcado elevado",
                                   f"Tasa de DROP {drop_rate:.1f}% según AST_VDADstats. "
                                   "Reducir nivel de marcado o aumentar agentes disponibles.", "#dc2626")
                    elif drop_rate > 3:
                        _vic_alert("DROP por encima del objetivo",
                                   f"Tasa de DROP {drop_rate:.1f}% (meta < 3%).", "#f59e0b")
        except Exception:
            pass

    if f_carrier is not None:
        st.info("📡 AST_carrier_log_report recibido — disponible para diagnóstico puntual de "
                "sobre-marcado por troncal (no se incluye en los 3 reportes diarios).")

    trend_lines = []
    for metric, label, fmt in [
        ("% Contactabilidad", "Contactabilidad", "{}"),
        ("% Cuelga en saludo", "Cuelga en saludo", "{}"),
        ("Promesas de pago (04+21)", "Promesas de pago", "{:.0f}"),
        ("Monto comprometido (S/)", "Monto comprometido (S/)", "{:.2f}"),
    ]:
        line = _vic_trend_line(hist_contact, metric, label, fmt)
        if line:
            trend_lines.append(line)

    if trend_lines:
        st.markdown("**Tendencia vs. día anterior:**")
        for line in trend_lines:
            st.markdown(f"- {line}")
    else:
        st.caption("Sube el Tablero_Contactabilidad_GRAL del día anterior para ver la tendencia.")

    points = _vic_critical(kp["total"], kp["humanos"], 0, kp["saludo"], None, None)
    color_map = {"rojo": "#dc2626", "naranja": "#f59e0b", "verde": "#22c55e"}
    for severity, title, body in points:
        if severity == "verde":
            st.success(f"✅ **{title}** — {body}")
        else:
            _vic_alert(title, body, color_map[severity])

    # ── Detalle de los 3 reportes ────────────────────────────────────────────
    st.markdown("#### Detalle de los Reportes")
    tab1, tab2, tab3 = st.tabs(["Tablero de Contactabilidad", "Control de Recontacto", "Tipificación de Gestión"])

    with tab1:
        _vic_resumen_cards(contact["resumen"])
        st.markdown("##### Por Estatus")
        _vic_estado_cards(contact["por_estado"], _VIC_GRAL_PROMESA, _VIC_GRAL_SALUDO)
        with st.expander("Ver tabla completa por estatus"):
            st.dataframe(contact["por_estado"], use_container_width=True, hide_index=True)
        if len(contact["por_entidad"]) > 0:
            st.markdown("##### Por Entidad (evasión y cuelga por estado)")
            st.dataframe(contact["por_entidad"], use_container_width=True, hide_index=True)
        if len(hist_contact) > 1:
            st.markdown("##### Histórico Acumulado")
            st.dataframe(hist_contact, use_container_width=True, hide_index=True)

    with tab2:
        if recont is None:
            st.info("No se detectó columna de teléfono / lead_id en el export para generar este reporte.")
        else:
            _vic_resumen_cards(recont["resumen"])
            st.markdown("##### Detalle por Lead")
            st.dataframe(recont["detalle"], use_container_width=True, hide_index=True)
            if hist_recont is not None and len(hist_recont) > 1:
                st.markdown("##### Histórico Acumulado")
                st.dataframe(hist_recont, use_container_width=True, hide_index=True)

    with tab3:
        _vic_resumen_cards(tipif["resumen"])
        st.markdown("##### Por Tipo de Gestión (Status)")
        st.dataframe(tipif["por_status"], use_container_width=True, hide_index=True)
        if len(tipif["por_estado_deudor"]) > 0:
            st.markdown("##### Por Estado del Deudor")
            st.dataframe(tipif["por_estado_deudor"], use_container_width=True, hide_index=True)
        if len(hist_tipif) > 1:
            st.markdown("##### Histórico Acumulado")
            st.dataframe(hist_tipif, use_container_width=True, hide_index=True)

    # ── Exportar los 3 reportes ──────────────────────────────────────────────
    st.divider()
    st.markdown("#### Descargar Reportes Actualizados")
    fecha_str = fecha.strftime("%Y%m%d")
    e1, e2, e3 = st.columns(3)

    with e1:
        buf1 = _vic_write_excel({
            "Resumen": contact["resumen"], "Por Estado": contact["por_estado"],
            "Por Entidad": contact["por_entidad"], "Histórico": hist_contact,
        })
        st.download_button("⬇️ Tablero_Contactabilidad_GRAL", data=buf1,
            file_name=f"Tablero_Contactabilidad_GRAL_{fecha_str}.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", type="primary")

    with e2:
        if recont is not None:
            buf2 = _vic_write_excel({
                "Resumen": recont["resumen"], "Detalle": recont["detalle"],
                "Histórico": hist_recont,
            })
            st.download_button("⬇️ Control_Recontacto_GRAL", data=buf2,
                file_name=f"Control_Recontacto_GRAL_{fecha_str}.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", type="primary")
        else:
            st.caption("No disponible (falta columna de teléfono/lead_id).")

    with e3:
        buf3 = _vic_write_excel({
            "Resumen": tipif["resumen"], "Por Status": tipif["por_status"],
            "Por Estado Deudor": tipif["por_estado_deudor"], "Histórico": hist_tipif,
        })
        st.download_button("⬇️ Tipificacion_Gestion_GRAL", data=buf3,
            file_name=f"Tipificacion_Gestion_GRAL_{fecha_str}.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", type="primary")

    return


def _coq_section(icon, title, subtitle=""):
    st.markdown(f"""
    <div style="margin:1.6rem 0 0.8rem 0;padding:0.9rem 1.2rem;border-radius:12px;
                background:linear-gradient(90deg,#1e2a4a 0%,#172033 100%);
                border-left:5px solid #3b82f6;">
        <div style="font-size:1.15rem;font-weight:800;color:#fff">{icon} {title}</div>
        {f'<div style="font-size:0.85rem;color:#8aa2c8;margin-top:2px">{subtitle}</div>' if subtitle else ''}
    </div>
    """, unsafe_allow_html=True)


def _coq_status_color(status, promesa, saludo_st):
    s = str(status).upper()
    if s in promesa:
        return "#22c55e"
    if s == str(saludo_st).upper() or s in {"AB", "DROP", "PDROP", "NA"}:
        return "#ef4444"
    if s in {"AA"}:
        return "#f59e0b"
    return "#3b82f6"


def page_coquimbo():
    st.markdown("""
    <div style="padding:1.4rem 1.6rem;border-radius:16px;margin-bottom:0.8rem;
                background:linear-gradient(120deg,#0ea5e9 0%,#2563eb 55%,#7c3aed 100%);">
        <div style="font-size:1.9rem;font-weight:800;color:#fff">📞 Reportes Diarios — Campaña Coquimbo</div>
        <div style="font-size:0.95rem;color:#e0e7ff;margin-top:4px">
            Sube los 3 archivos del día y genera el análisis completo: contactabilidad, recontacto y
            tipificación de gestión, con gráficas y tendencia día a día.
        </div>
    </div>
    """, unsafe_allow_html=True)

    with st.expander("ℹ️ Reglas aplicadas en el cálculo"):
        st.markdown("""
- **Promesa de pago** = status `1B` + `1O` (siempre se suman ambos).
- **Status `1L`** = cuelga en saludo (rechazo temprano), no cuenta como gestión.
- **Humanos** = status `2, 4, 6, 1B, 1C, 1D, 1E, 1F, 1G, 1H, 1I, 1J, 1K, 1L, 1N, 1O, 2B, 2C, 2D, 2E, 2F, 3C, 3D, 3E`
  → hoja *Por Entidad* (evasión y cuelga por estado, agrupado por lista/cartera).
- **Sábado** = media jornada (8 AM–12 PM); se compara sábado contra sábado.
- El día nuevo se agrega al **acumulado (Histórico)** de cada reporte, leyendo el tablero del día anterior.
- **Monto comprometido** sale del campo `last_name` (campo reutilizado del export). El campo `first_name`
  trae el ID del deudor y no se usa para tipificación.
        """)

    _coq_section("📁", "1. Archivos frescos del día", "3 archivos de la jornada")
    f_amd     = st.file_uploader("AST_AMD_log_report (.csv)", type=None, key="coq_amd")
    f_vdad    = st.file_uploader("Estatus de llamadas / AST_VDADstats (.csv)", type=None, key="coq_vdad")
    f_export  = st.file_uploader(
        "Estados VICIdial — EXPORT_CALL_REPORT (.txt/.csv)  ⭐ insumo principal",
        type=None, key="coq_export")

    _coq_section("🗄️", "2. Tableros del día anterior", "3 archivos del acumulado histórico")
    f_contact_prev    = st.file_uploader("Tablero_Contactabilidad_Coquimbo (.xlsx)", type=None, key="coq_contact_prev")
    f_recontacto_prev = st.file_uploader("Control_Recontacto_Coquimbo (.xlsx)", type=None, key="coq_recontacto_prev")
    f_tipif_prev      = st.file_uploader("Tipificacion_Gestion_Coquimbo (.xlsx)", type=None, key="coq_tipif_prev")

    st.divider()

    if f_export is None:
        st.info("Sube al menos el **export de Estados VICIdial** para generar los 3 reportes.")
        return

    if not st.button("🚀 Generar Reportes", type="primary"):
        return

    with st.spinner("Procesando..."):
        try:
            df, legend = _vic_load_with_legend(f_export)
        except Exception as e:
            st.error(f"No se pudo leer el export: {e}")
            return

        c = _vic_export_cols(df, monto_subs=("last_name",), estado_subs=("first_name",))
        if not c["status"]:
            st.error("No se encontró la columna **status** en el export. "
                     "Verifica que el archivo incluya esa columna.")
            return
        df["_st"] = _vic_status_norm(df[c["status"]])

        if f_vdad is not None:
            try:
                _, legend_vdad = _vic_load_with_legend(f_vdad)
                if not legend_vdad:
                    vdf = _vic_load(f_vdad)
                    col_cod  = _vic_find(vdf, "codigo", "código", "status", "code")
                    col_desc = _vic_find(vdf, "descripcion", "descripción", "desc")
                    if col_cod and col_desc:
                        legend_vdad = {
                            str(r[col_cod]).strip().upper(): str(r[col_desc]).strip()
                            for _, r in vdf.iterrows()
                            if str(r[col_cod]).strip() and str(r[col_cod]).strip().upper() != "NAN"
                        }
                legend = {**legend_vdad, **legend}
            except Exception:
                pass

        labels_full = {**legend, **_COQ_LABELS}

        fecha = _vic_extract_date_from_name(f_export.name) or date.today()
        df, es_sabado = _vic_filter_jornada(df, c, fecha)
        if es_sabado:
            st.info("📅 Detectado **sábado** — se filtró la jornada a 8:00–12:00 (media jornada). "
                    "Compara este reporte contra el sábado anterior en la hoja *Histórico*.")

        contact = _vic_tablero_contactabilidad(df, c, fecha, promesa=_COQ_PROMESA,
                                                saludo_st=_COQ_SALUDO, humanos_st=_COQ_HUMANOS, labels=labels_full)
        recont  = _vic_control_recontacto(df, c, fecha, promesa=_COQ_PROMESA, labels=labels_full)
        tipif   = _vic_tipificacion_gestion(df, c, fecha, labels=labels_full)

        hist_contact = _vic_append_historico(f_contact_prev, contact["resumen"])
        hist_tipif   = _vic_append_historico(f_tipif_prev, tipif["resumen"])
        hist_recont  = _vic_append_historico(f_recontacto_prev, recont["resumen"]) if recont else None

    # ── KPIs ──────────────────────────────────────────────────────────────────
    _coq_section("📊", "Resumen del Día", f"Jornada del {fecha.strftime('%d/%m/%Y')}"
                  + (" · sábado, media jornada" if es_sabado else ""))
    k1, k2, k3, k4, k5 = st.columns(5)
    kp = contact["kpis"]
    cr_color = "#27ae60" if kp["humanos"]/kp["total"] >= 0.35 else "#f39c12" if kp["humanos"]/kp["total"] >= 0.20 else "#e74c3c"
    _vic_kpi(k1, "Total llamadas", f"{kp['total']:,}")
    _vic_kpi(k2, "Contactabilidad (humanos)", _vic_pct(kp["humanos"], kp["total"]), color=cr_color)
    _vic_kpi(k3, "Promesas de pago (1B+1O)", f"{kp['promesas']:,}", color="#22c55e")
    _vic_kpi(k4, "Monto comprometido", f"$ {kp['monto']:,.2f}", color="#a855f7")
    sal_color = "#e74c3c" if kp["saludo"]/kp["total"] > 0.40 else "#f39c12" if kp["saludo"]/kp["total"] > 0.25 else "#27ae60"
    _vic_kpi(k5, "Cuelga en saludo (1L)", _vic_pct(kp["saludo"], kp["total"]), color=sal_color)

    # ── Gráficas: composición + por estado ──────────────────────────────────
    _coq_section("📈", "Análisis Visual", "Composición de la jornada y top de disposiciones")
    g1, g2 = st.columns([1, 2])

    with g1:
        no_contact = max(kp["total"] - kp["humanos"], 0)
        fig_donut = go.Figure(data=[go.Pie(
            labels=["Contacto humano", "No contactable / máquina"],
            values=[kp["humanos"], no_contact],
            hole=0.62,
            marker=dict(colors=["#22c55e", "#ef4444"]),
            textinfo="percent",
            textfont=dict(color="#fff", size=14),
        )])
        fig_donut.update_layout(
            title=dict(text="Contactabilidad", font=dict(color="#e2e8f0", size=15)),
            paper_bgcolor="rgba(0,0,0,0)", plot_bgcolor="rgba(0,0,0,0)",
            showlegend=True, legend=dict(font=dict(color="#cbd5e1"), orientation="h", y=-0.1),
            margin=dict(t=40, b=10, l=10, r=10), height=320,
            annotations=[dict(text=f"{_vic_pct(kp['humanos'], kp['total'])}", x=0.5, y=0.5,
                               font=dict(size=22, color="#fff"), showarrow=False)],
        )
        st.plotly_chart(fig_donut, use_container_width=True)

    with g2:
        top_estado = contact["por_estado"].sort_values("Llamadas", ascending=False).head(15).copy()
        top_estado["_label"] = top_estado.apply(
            lambda r: (r["Descripción"][:28] if r["Descripción"] and r["Descripción"] != "—" else r["Status"]),
            axis=1)
        colors_bar = [_coq_status_color(s, _COQ_PROMESA, _COQ_SALUDO) for s in top_estado["Status"]]
        fig_bar = go.Figure(go.Bar(
            x=top_estado["Llamadas"], y=top_estado["_label"], orientation="h",
            marker_color=colors_bar, text=top_estado["Llamadas"], textposition="outside",
            customdata=top_estado["Status"],
            hovertemplate="<b>%{customdata}</b> — %{y}<br>Llamadas: %{x}<extra></extra>",
        ))
        fig_bar.update_layout(
            title=dict(text="Top disposiciones del día", font=dict(color="#e2e8f0", size=15)),
            paper_bgcolor="rgba(0,0,0,0)", plot_bgcolor="rgba(0,0,0,0)",
            font=dict(color="#cbd5e1"), height=320,
            yaxis=dict(autorange="reversed", gridcolor="#27314a"),
            xaxis=dict(gridcolor="#27314a"),
            margin=dict(t=40, b=10, l=10, r=10),
        )
        st.plotly_chart(fig_bar, use_container_width=True)
        st.caption("🟢 Promesa de pago · 🔴 Problema operativo (saludo / agente no disponible / drop) · 🟠 Buzón · 🔵 Otro")

    # ── Alertas / tendencia ──────────────────────────────────────────────────
    _coq_section("🚨", "Alertas y Tendencia", "Comparación contra el día anterior y focos rojos")
    saludo_rate = kp["saludo"] / kp["total"] if kp["total"] else 0
    if saludo_rate > 0.40:
        _vic_alert("Evasión / cuelga en saludo CRÍTICA",
                   f"{saludo_rate*100:.1f}% de las llamadas cuelgan en el saludo (status 1L). "
                   "Revisar guion de apertura, horario de marcado y calidad de la base.", "#dc2626")
    elif saludo_rate > 0.25:
        _vic_alert("Cuelga en saludo elevado",
                   f"{saludo_rate*100:.1f}% cuelgan en el saludo. Vigilar tendencia.", "#f59e0b")

    if f_vdad is not None:
        try:
            vdf = _vic_load(f_vdad)
            drop_col  = _vic_find(vdf, "drop")
            calls_col = _vic_find(vdf, "call")
            if drop_col and calls_col:
                drops = pd.to_numeric(vdf[drop_col], errors="coerce").fillna(0).sum()
                calls = pd.to_numeric(vdf[calls_col], errors="coerce").fillna(0).sum()
                if calls > 0:
                    drop_rate = drops / calls * 100
                    if drop_rate > 5:
                        _vic_alert("DROP / sobre-marcado elevado",
                                   f"Tasa de DROP {drop_rate:.1f}% según el reporte de estatus de llamadas. "
                                   "Reducir nivel de marcado o aumentar agentes disponibles.", "#dc2626")
                    elif drop_rate > 3:
                        _vic_alert("DROP por encima del objetivo",
                                   f"Tasa de DROP {drop_rate:.1f}% (meta < 3%).", "#f59e0b")
        except Exception:
            pass

    trend_lines = []
    for metric, label, fmt in [
        ("% Contactabilidad", "Contactabilidad", "{}"),
        ("% Cuelga en saludo", "Cuelga en saludo", "{}"),
        ("Promesas de pago (1B+1O)", "Promesas de pago", "{:.0f}"),
        ("Monto comprometido (S/)", "Monto comprometido ($)", "{:.2f}"),
    ]:
        line = _vic_trend_line(hist_contact, metric, label, fmt)
        if line:
            trend_lines.append(line)

    if trend_lines:
        st.markdown("**Tendencia vs. día anterior:**")
        for line in trend_lines:
            st.markdown(f"- {line}")
    else:
        st.caption("Sube el Tablero_Contactabilidad_Coquimbo del día anterior para ver la tendencia.")

    # ── Detalle de los 3 reportes ────────────────────────────────────────────
    _coq_section("🗂️", "Detalle de los Reportes", "Tablero, recontacto y tipificación completos")
    tab1, tab2, tab3 = st.tabs(["📊 Tablero de Contactabilidad", "🔁 Control de Recontacto", "🏷️ Tipificación de Gestión"])

    with tab1:
        st.markdown("##### Resumen")
        _vic_resumen_cards(contact["resumen"])
        st.markdown("##### Por Estatus")
        _vic_estado_cards(contact["por_estado"], _COQ_PROMESA, _COQ_SALUDO)
        with st.expander("Ver tabla completa por estatus"):
            st.dataframe(contact["por_estado"], use_container_width=True, hide_index=True)
        if len(contact["por_entidad"]) > 0:
            st.markdown("##### Por Entidad (evasión y cuelga por estado)")
            ent_col = contact["por_entidad"].columns[0]
            _vic_entidad_cards(contact["por_entidad"], ent_col)
            with st.expander("Ver tabla completa por entidad"):
                st.dataframe(contact["por_entidad"], use_container_width=True, hide_index=True)
        if len(hist_contact) > 1:
            st.markdown("##### Histórico Acumulado")
            st.dataframe(hist_contact, use_container_width=True, hide_index=True)
            if "% Contactabilidad" in hist_contact.columns:
                hc = hist_contact.copy()
                hc["_pct"] = pd.to_numeric(hc["% Contactabilidad"].astype(str).str.replace("%", ""), errors="coerce")
                fig_trend = go.Figure(go.Scatter(
                    x=hc["Fecha"], y=hc["_pct"], mode="lines+markers",
                    line=dict(color="#3b82f6", width=3), marker=dict(size=8, color="#60a5fa"),
                    fill="tozeroy", fillcolor="rgba(59,130,246,0.15)",
                ))
                fig_trend.update_layout(
                    title=dict(text="Tendencia de contactabilidad", font=dict(color="#e2e8f0", size=15)),
                    paper_bgcolor="rgba(0,0,0,0)", plot_bgcolor="rgba(0,0,0,0)",
                    font=dict(color="#cbd5e1"), height=280,
                    xaxis=dict(gridcolor="#27314a"), yaxis=dict(gridcolor="#27314a", title="%"),
                    margin=dict(t=40, b=10, l=10, r=10),
                )
                st.plotly_chart(fig_trend, use_container_width=True)

    with tab2:
        if recont is None:
            st.info("No se detectó columna de teléfono / lead_id en el export para generar este reporte.")
        else:
            _vic_resumen_cards(recont["resumen"])
            st.markdown("##### Detalle por Lead")
            st.dataframe(recont["detalle"], use_container_width=True, hide_index=True)
            if hist_recont is not None and len(hist_recont) > 1:
                st.markdown("##### Histórico Acumulado")
                st.dataframe(hist_recont, use_container_width=True, hide_index=True)

    with tab3:
        st.markdown("##### Resumen")
        _vic_resumen_cards(tipif["resumen"])
        st.markdown("##### Por Tipo de Gestión (Status)")
        ts = tipif["por_status"].sort_values("llamadas", ascending=False)
        fig_tip = go.Figure(go.Bar(
            x=ts["Status"], y=ts["llamadas"],
            marker_color=[_coq_status_color(s, _COQ_PROMESA, _COQ_SALUDO) for s in ts["Status"]],
            text=ts["llamadas"], textposition="outside",
            customdata=ts["Descripción"],
            hovertemplate="<b>%{x}</b> — %{customdata}<br>Llamadas: %{y}<extra></extra>",
        ))
        fig_tip.update_layout(
            title=dict(text="Distribución de la gestión por status", font=dict(color="#e2e8f0", size=15)),
            paper_bgcolor="rgba(0,0,0,0)", plot_bgcolor="rgba(0,0,0,0)",
            font=dict(color="#cbd5e1"), height=340,
            xaxis=dict(gridcolor="#27314a"), yaxis=dict(gridcolor="#27314a"),
            margin=dict(t=40, b=10, l=10, r=10),
        )
        st.plotly_chart(fig_tip, use_container_width=True)
        st.dataframe(tipif["por_status"], use_container_width=True, hide_index=True)
        if len(tipif["por_estado_deudor"]) > 0:
            st.markdown("##### Por ID de Deudor (referencia)")
            st.dataframe(tipif["por_estado_deudor"], use_container_width=True, hide_index=True)
        if len(hist_tipif) > 1:
            st.markdown("##### Histórico Acumulado")
            st.dataframe(hist_tipif, use_container_width=True, hide_index=True)

    # ── Exportar los 3 reportes ──────────────────────────────────────────────
    st.divider()
    _coq_section("⬇️", "Descargar Reportes Actualizados", "Listos para enviar o archivar en el histórico")
    fecha_str = fecha.strftime("%Y%m%d")
    e1, e2, e3 = st.columns(3)

    with e1:
        buf1 = _vic_write_excel({
            "Resumen": contact["resumen"], "Por Estado": contact["por_estado"],
            "Por Entidad": contact["por_entidad"], "Histórico": hist_contact,
        })
        st.download_button("⬇️ Tablero_Contactabilidad_Coquimbo", data=buf1,
            file_name=f"Tablero_Contactabilidad_Coquimbo_{fecha_str}.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", type="primary")

    with e2:
        if recont is not None:
            buf2 = _vic_write_excel({
                "Resumen": recont["resumen"], "Detalle": recont["detalle"],
                "Histórico": hist_recont,
            })
            st.download_button("⬇️ Control_Recontacto_Coquimbo", data=buf2,
                file_name=f"Control_Recontacto_Coquimbo_{fecha_str}.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", type="primary")
        else:
            st.caption("No disponible (falta columna de teléfono/lead_id).")

    with e3:
        buf3 = _vic_write_excel({
            "Resumen": tipif["resumen"], "Por Status": tipif["por_status"],
            "Por ID Deudor": tipif["por_estado_deudor"], "Histórico": hist_tipif,
        })
        st.download_button("⬇️ Tipificacion_Gestion_Coquimbo", data=buf3,
            file_name=f"Tipificacion_Gestion_Coquimbo_{fecha_str}.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", type="primary")

    return


# ─────────────────────────────────────────────────────────────────────────────
# PAGINA: REPORTE DE CAMPAÑA (resumen, agentes, promesas, alertas, plan)
# ─────────────────────────────────────────────────────────────────────────────

def _norm_txt(s) -> str:
    s = str(s).upper().strip()
    s = unicodedata.normalize("NFKD", s)
    return "".join(c for c in s if not unicodedata.combining(c))


def _camp_categoria(status_norm: str) -> str:
    s = status_norm
    if "ADMINISTRATIVA" in s:
        return "administrativa"
    if "AGENTE NO DISPONIBLE" in s:
        return "agente_no_disponible"
    if "PRE-ROUTING DROP" in s or "PRE ROUTING DROP" in s or "PREROUTING" in s:
        return "drop"
    if "OCUPADO" in s:
        return "ocupado"
    if "NO ANSWER" in s and "AUTODIAL" in s:
        return "no_answer_auto"
    if "NO ANSWER" in s or "NO CONTESTA" in s:
        return "no_contesta"
    if "BUZON" in s and ("AUTOMATIC" in s or "AUTODIAL" in s):
        return "buzon_automatico"
    if "BUZON" in s:
        return "directo_buzon"
    if "CUELGA" in s:
        return "cuelga"
    if "INCUMPLIDA" in s:
        return "promesa_incumplida"
    if "PROMESA" in s and "PAGO" in s:
        return "promesa_pago"
    if "NEGATIVA" in s:
        return "negativa_pago"
    return "contacto_util"


_CAMP_NO_CONTACTABLE = {"ocupado", "no_answer_auto", "no_contesta", "buzon_automatico"}
_CAMP_BUZON          = {"buzon_automatico", "directo_buzon"}
_CAMP_PROBLEMA       = {"ocupado", "no_answer_auto", "no_contesta", "buzon_automatico",
                         "agente_no_disponible", "cuelga", "drop"}
_CAMP_EXCLUIR_AGENTE = {"administrativa", "ocupado", "no_answer_auto", "buzon_automatico", "drop"}

_RE_CAMP_MONTO = re.compile(r"POR\s*\$?\s*([\d,]+(?:\.\d+)?)", re.IGNORECASE)
_RE_CAMP_FECHA = re.compile(r"PARA\s+EL\s+(\d{1,2}[\.\-/]\d{1,2}[\.\-/]\d{2,4})", re.IGNORECASE)


def _camp_resumen(df: pd.DataFrame) -> dict | None:
    disp_col  = _vic_find(df, "disposic", "status", "estado")
    calls_col = _vic_find(df, "llamada", "calls", "cantidad")
    if not disp_col or not calls_col:
        return None
    out = df[[disp_col, calls_col]].copy()
    out.columns = ["Disposición", "Llamadas"]
    out["Llamadas"] = pd.to_numeric(out["Llamadas"], errors="coerce").fillna(0)
    out["_cat"] = out["Disposición"].apply(lambda s: _camp_categoria(_norm_txt(s)))
    total = out["Llamadas"].sum()
    no_contactable = out.loc[out["_cat"].isin(_CAMP_NO_CONTACTABLE), "Llamadas"].sum()
    buzon_total     = out.loc[out["_cat"].isin(_CAMP_BUZON), "Llamadas"].sum()
    administrativa  = out.loc[out["_cat"] == "administrativa", "Llamadas"].sum()
    contacto_util   = total - no_contactable - administrativa
    ocupado         = out.loc[out["_cat"] == "ocupado", "Llamadas"].sum()
    agente_no_disp  = out.loc[out["_cat"] == "agente_no_disponible", "Llamadas"].sum()
    return {
        "tabla": out.sort_values("Llamadas", ascending=False),
        "total": total, "no_contactable": no_contactable, "buzon_total": buzon_total,
        "contacto_util": contacto_util, "ocupado": ocupado, "agente_no_disponible": agente_no_disp,
    }


def _camp_detalle_cols(df: pd.DataFrame) -> dict:
    return {
        "date":     _vic_find(df, "call_date", "fecha", "date"),
        "agent":    _vic_find(df, "full_name", "agente", "user", "nombre"),
        "status":   _vic_find(df, "status_name", "status", "disposic"),
        "duration": _vic_find(df, "length_in_sec", "duration", "duracion"),
        "comments": _vic_find(df, "comments", "comentario"),
        "phone":    _vic_find(df, "phone_number_dialed", "phone", "telefono"),
        "lead_id":  _vic_find(df, "lead_id"),
    }


def _camp_agentes(df: pd.DataFrame, cols: dict) -> pd.DataFrame:
    d = df.copy()
    d["_st"]  = d[cols["status"]].apply(lambda s: _norm_txt(s))
    d["_cat"] = d["_st"].apply(_camp_categoria)
    d = d[~d["_cat"].isin(_CAMP_EXCLUIR_AGENTE)]
    if cols["duration"]:
        d["_dur"] = pd.to_numeric(d[cols["duration"]], errors="coerce").fillna(0)

    rows = []
    for agente, sub in d.groupby(cols["agent"]):
        total         = len(sub)
        cuelgues      = int((sub["_cat"] == "cuelga").sum())
        directo_buzon = int((sub["_cat"] == "directo_buzon").sum())
        no_contestan  = int((sub["_cat"] == "no_contesta").sum())
        negativas     = int((sub["_cat"] == "negativa_pago").sum())
        promesas      = int((sub["_cat"] == "promesa_pago").sum())
        incumplidas   = int((sub["_cat"] == "promesa_incumplida").sum())
        contacto_util = negativas + promesas + incumplidas + int((sub["_cat"] == "contacto_util").sum())
        tasa = contacto_util / total * 100 if total else 0
        avg_dur = sub["_dur"].mean() if cols["duration"] else None
        rows.append({
            "Ejecutivo": agente, "Total": total, "Tasa Contacto Util %": round(tasa, 1),
            "Cuelgues": cuelgues, "Directo Buzon": directo_buzon, "No Contestan": no_contestan,
            "Negativas Pago": negativas, "Promesas Pago": promesas, "Promesas Incumplidas": incumplidas,
            "Duracion Prom (seg)": round(avg_dur, 0) if avg_dur is not None else None,
        })
    return pd.DataFrame(rows).sort_values("Tasa Contacto Util %", ascending=False).reset_index(drop=True)


def _camp_promesas(df: pd.DataFrame, cols: dict):
    if not cols["comments"] or not cols["status"]:
        return None, None
    d = df.copy()
    d["_st"]  = d[cols["status"]].apply(lambda s: _norm_txt(s))
    d["_cat"] = d["_st"].apply(_camp_categoria)

    def extract(row):
        c = str(row[cols["comments"]]) if pd.notna(row[cols["comments"]]) else ""
        m = _RE_CAMP_MONTO.search(c)
        f = _RE_CAMP_FECHA.search(c)
        monto = float(m.group(1).replace(",", "")) if m else None
        fecha = f.group(1) if f else None
        return pd.Series({"_monto": monto, "_fecha_compromiso": fecha})

    def build(cat):
        sub = d[d["_cat"] == cat].copy()
        if len(sub) == 0:
            return pd.DataFrame(columns=["Hora", "Ejecutivo", "Telefono", "Lead ID", "Monto",
                                          "Fecha Compromiso", "Comentario"])
        extras = sub.apply(extract, axis=1)
        sub = pd.concat([sub.reset_index(drop=True), extras.reset_index(drop=True)], axis=1)
        return pd.DataFrame({
            "Hora":             sub[cols["date"]] if cols["date"] else "",
            "Ejecutivo":        sub[cols["agent"]] if cols["agent"] else "",
            "Telefono":         sub[cols["phone"]] if cols["phone"] else "",
            "Lead ID":          sub[cols["lead_id"]] if cols["lead_id"] else "",
            "Monto":            sub["_monto"],
            "Fecha Compromiso": sub["_fecha_compromiso"],
            "Comentario":       sub[cols["comments"]],
        })

    return build("promesa_pago"), build("promesa_incumplida")


def _camp_export_promesas(df_prom: pd.DataFrame | None, df_incump: pd.DataFrame | None) -> io.BytesIO:
    from openpyxl.styles import Font, PatternFill, Alignment
    from openpyxl.utils import get_column_letter

    buf = io.BytesIO()
    with pd.ExcelWriter(buf, engine="openpyxl") as writer:
        for name, sdf in [("Promesas de Pago", df_prom), ("Promesas Incumplidas", df_incump)]:
            sdf = sdf if sdf is not None else pd.DataFrame()
            sdf.to_excel(writer, sheet_name=name[:31], index=False)
            ws = writer.sheets[name[:31]]
            if len(sdf) == 0:
                continue
            header_fill = PatternFill(start_color="1E2535", end_color="1E2535", fill_type="solid")
            header_font = Font(color="FFFFFF", bold=True)
            for col_idx, col_name in enumerate(sdf.columns, start=1):
                cell = ws.cell(row=1, column=col_idx)
                cell.fill = header_fill
                cell.font = header_font
                cell.alignment = Alignment(horizontal="center")
                ws.column_dimensions[get_column_letter(col_idx)].width = max(14, len(str(col_name)) + 2)
            if "Monto" in sdf.columns:
                col_idx = list(sdf.columns).index("Monto") + 1
                col_letter = get_column_letter(col_idx)
                n = len(sdf)
                for r in range(2, n + 2):
                    ws.cell(row=r, column=col_idx).number_format = '"$"#,##0.00'
                total_row = n + 2
                label_cell = ws.cell(row=total_row, column=col_idx - 1, value="TOTAL")
                label_cell.font = Font(bold=True)
                sum_cell = ws.cell(row=total_row, column=col_idx, value=f"=SUM({col_letter}2:{col_letter}{n+1})")
                sum_cell.number_format = '"$"#,##0.00'
                sum_cell.font = Font(bold=True)
    buf.seek(0)
    return buf


def _camp_alertas(resumen, agentes_df, df_prom, df_incump):
    alerts = []

    if resumen is not None and resumen["total"] > 0:
        ocupado_pct = resumen["ocupado"] / resumen["total"] * 100
        if ocupado_pct > 35:
            alerts.append(("rojo", "Ocupado automático elevado",
                f"{ocupado_pct:.1f}% de las llamadas son 'Ocupado automático' (umbral 35%). "
                "Revisar calidad de la base de datos."))
        if resumen["agente_no_disponible"] > 100:
            alerts.append(("naranja", "Agente no disponible elevado",
                f"{int(resumen['agente_no_disponible'])} casos de 'Agente no disponible' (umbral 100). "
                "Revisar dial ratio / staffing."))

    if agentes_df is not None and len(agentes_df) > 0:
        for _, r in agentes_df[agentes_df["Tasa Contacto Util %"] == 0].iterrows():
            alerts.append(("rojo", f"{r['Ejecutivo']}: 0% de contacto útil",
                "Revisión individual urgente — sin gestiones efectivas en el día."))
        cuelga_pct = agentes_df["Cuelgues"] / agentes_df["Total"].replace(0, np.nan) * 100
        for (_, r), pct in zip(agentes_df.iterrows(), cuelga_pct):
            if pct and pct > 45:
                alerts.append(("naranja", f"{r['Ejecutivo']}: cuelgan llamada > 45%",
                    f"{pct:.1f}% de sus llamadas cuelgan. Revisar script de apertura."))

    if df_incump is not None and len(df_incump) > 0 and "Monto" in df_incump.columns:
        total_incump = pd.to_numeric(df_incump["Monto"], errors="coerce").fillna(0).sum()
        if total_incump > 3000:
            alerts.append(("rojo", "Promesas incumplidas superan $3,000",
                f"Total en riesgo: ${total_incump:,.2f}. Requiere seguimiento urgente."))

    hoy_variants = {datetime.now().strftime("%d/%m/%y"), datetime.now().strftime("%d/%m/%Y"),
                    datetime.now().strftime("%d.%m.%y"), datetime.now().strftime("%d-%m-%y")}
    hoy_norm = {v.replace(".", "/").replace("-", "/") for v in hoy_variants}
    for label, df_ in [("Promesa de pago", df_prom), ("Promesa incumplida", df_incump)]:
        if df_ is not None and len(df_) > 0 and "Fecha Compromiso" in df_.columns:
            fechas_norm = df_["Fecha Compromiso"].astype(str).str.replace(".", "/", regex=False) \
                                                  .str.replace("-", "/", regex=False)
            n = fechas_norm.isin(hoy_norm).sum()
            if n > 0:
                alerts.append(("naranja", f"{n} '{label}' vencen HOY",
                    "Contactar al cliente hoy para confirmar el pago."))

    return alerts


def _camp_plan(alerts, agentes_df) -> dict:
    plan = {"alta": [], "media": [], "baja": []}
    for sev, title, body in alerts:
        (plan["alta"] if sev == "rojo" else plan["media"]).append(f"**{title}** — {body}")
    if agentes_df is not None and len(agentes_df) > 0:
        bajos = agentes_df[agentes_df["Tasa Contacto Util %"] < 25]
        if len(bajos) > 0:
            nombres = ", ".join(bajos["Ejecutivo"].astype(str).head(5))
            plan["media"].append(f"Capacitación / seguimiento para: {nombres} (tasa de contacto útil < 25%).")
    plan["baja"].append("Revisar listas 'New Lead' sin gestionar y priorizarlas en la próxima jornada.")
    plan["baja"].append("Si 'Ocupado automático' sigue alto, ajustar horarios de marcación y dial ratio (AMD).")
    return plan


def page_campana():
    st.markdown("## 📊 Reporte Diario de Campaña")
    st.markdown("*Resumen de campaña, rendimiento por ejecutivo, promesas de pago, alertas y plan de acción.*")
    st.divider()

    with st.expander("ℹ️ Cómo se clasifican las disposiciones"):
        st.markdown("""
- **No contactables**: Ocupado automático, Buzón automático, No Answer / No contesta.
- **Buzón total**: Buzón automático + Directo a buzón.
- **Contacto útil**: Negativa de pago, Promesa de pago, Promesa incumplida y otras gestiones con conversación real.
- Se **excluyen** del análisis por ejecutivo: Llamada administrativa, Ocupado automático,
  Buzón automático, No Answer AutoDial y Outbound Pre-Routing Drop.
        """)

    c1, c2 = st.columns(2)
    with c1:
        f_resumen = st.file_uploader("Reporte de Estados de Campaña (Disposición, Llamadas, % )",
                                      type=None, key="camp_resumen")
    with c2:
        f_detalle = st.file_uploader("Reporte Detallado de Llamadas (call_date, full_name, status_name, comments...)",
                                      type=None, key="camp_detalle")

    if f_resumen is None and f_detalle is None:
        st.info("Sube al menos uno de los dos reportes para comenzar.")
        return

    resumen = None
    agentes_df = None
    df_prom = df_incump = None

    if f_resumen is not None:
        try:
            rdf = _vic_load(f_resumen)
            resumen = _camp_resumen(rdf)
            if resumen is None:
                st.warning("No se reconocieron las columnas de Disposición / Llamadas en el reporte de estados.")
        except Exception as e:
            st.warning(f"No se pudo leer el reporte de estados: {e}")

    if f_detalle is not None:
        try:
            ddf = _vic_load(f_detalle)
            cols = _camp_detalle_cols(ddf)
            if not cols["status"] or not cols["agent"]:
                st.warning("No se encontraron las columnas 'status_name' / 'full_name' en el reporte detallado.")
            else:
                agentes_df = _camp_agentes(ddf, cols)
                df_prom, df_incump = _camp_promesas(ddf, cols)
        except Exception as e:
            st.warning(f"No se pudo leer el reporte detallado: {e}")

    # ── Módulo 1: Resumen de campaña ─────────────────────────────────────────
    if resumen is not None:
        st.markdown("#### 1. Resumen de Campaña")
        k1, k2, k3, k4 = st.columns(4)
        _vic_kpi(k1, "Total llamadas", f"{int(resumen['total']):,}")
        _vic_kpi(k2, "% No contactables", _vic_pct(resumen["no_contactable"], resumen["total"]), color="#f59e0b")
        _vic_kpi(k3, "% Contacto útil", _vic_pct(resumen["contacto_util"], resumen["total"]), color="#22c55e")
        _vic_kpi(k4, "% Buzón total", _vic_pct(resumen["buzon_total"], resumen["total"]), color="#8b5cf6")

        tabla = resumen["tabla"]
        colors = ["#ef4444" if c in _CAMP_PROBLEMA else "#3b82f6" for c in tabla["_cat"]]
        fig = go.Figure(go.Bar(x=tabla["Disposición"], y=tabla["Llamadas"], marker_color=colors))
        fig.update_layout(paper_bgcolor="#1e2535", plot_bgcolor="#1e2535",
                          font=dict(color="#aaa", size=11), margin=dict(t=30, b=80, l=10, r=10),
                          title="Distribución de Disposiciones (rojo = problema operativo)",
                          xaxis=dict(tickangle=-45))
        st.plotly_chart(fig, use_container_width=True)
        st.divider()

    # ── Módulo 2: Rendimiento por ejecutivo ──────────────────────────────────
    if agentes_df is not None:
        st.markdown("#### 2. Rendimiento por Ejecutivo")

        def _color_tasa(val):
            if val >= 50:
                return "background-color:#16341f;color:#4ade80"
            if val >= 25:
                return "background-color:#3a2a00;color:#fbbf24"
            return "background-color:#3a0f0f;color:#f87171"

        styled = agentes_df.style.map(_color_tasa, subset=["Tasa Contacto Util %"])
        st.dataframe(styled, use_container_width=True, hide_index=True)
        st.divider()

    # ── Módulo 3: Promesas de pago ───────────────────────────────────────────
    if df_prom is not None or df_incump is not None:
        st.markdown("#### 3. Promesas de Pago y Seguimiento")
        t1, t2 = st.tabs(["Promesas de Pago", "Promesas Incumplidas"])
        with t1:
            if df_prom is not None and len(df_prom) > 0:
                st.dataframe(df_prom, use_container_width=True, hide_index=True)
            else:
                st.info("No se detectaron promesas de pago.")
        with t2:
            if df_incump is not None and len(df_incump) > 0:
                total_incump = pd.to_numeric(df_incump["Monto"], errors="coerce").fillna(0).sum()
                st.metric("Monto total en riesgo", f"${total_incump:,.2f}")
                st.dataframe(df_incump, use_container_width=True, hide_index=True)
            else:
                st.info("No se detectaron promesas incumplidas.")

        buf = _camp_export_promesas(df_prom, df_incump)
        st.download_button("⬇️ Descargar Promesas (Excel)", data=buf,
            file_name=f"promesas_{datetime.now().strftime('%Y%m%d')}.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", type="primary")
        st.divider()

    # ── Módulo 4: Alertas automáticas ────────────────────────────────────────
    st.markdown("#### 4. Alertas Automáticas")
    alerts = _camp_alertas(resumen, agentes_df, df_prom, df_incump)
    if not alerts:
        st.success("✅ Sin alertas — no se detectaron patrones críticos.")
    else:
        color_map = {"rojo": "#dc2626", "naranja": "#f59e0b"}
        for sev, title, body in alerts:
            _vic_alert(title, body, color_map[sev])
    st.divider()

    # ── Módulo 5: Plan de acción ─────────────────────────────────────────────
    st.markdown("#### 5. Plan de Acción")
    plan = _camp_plan(alerts, agentes_df)
    for nivel, label in [("alta", "Prioridad Alta"), ("media", "Prioridad Media"), ("baja", "Prioridad Baja")]:
        if plan[nivel]:
            st.markdown(f"**{label}**")
            for item in plan[nivel]:
                st.markdown(f"- {item}")


# ─────────────────────────────────────────────────────────────────────────────
# ROUTER PRINCIPAL
# ─────────────────────────────────────────────────────────────────────────────

def _get_or_create_empresa(nombre: str, slug: str) -> int:
    """Retorna el id de la empresa, creándola si no existe."""
    try:
        from database import get_empresa_id_by_slug
        eid = get_empresa_id_by_slug(slug)
        if eid:
            return eid
        create_empresa(nombre, slug)
        return get_empresa_id_by_slug(slug) or 1
    except Exception:
        return 1


def _bootstrap_users():
    """Crea empresas y usuarios por defecto al primer inicio."""
    try:
        # ── Empresa 1: Cuzco ──────────────────────────────────────────────────
        eid_cuzco = _get_or_create_empresa("Cuzco", "cuzco")
        existing_cuzco = set(get_all_users(eid_cuzco)["username"].tolist()) if len(get_all_users(eid_cuzco)) > 0 else set()
        for username, password, nombre, rol in [
            ("sup1", "Cuzco@Sup1",  "Sup 1", "admin"),
            ("sup2", "Cuzco@Sup2",  "Sup 2", "admin"),
        ]:
            if username not in existing_cuzco:
                create_user(username, password, nombre, f"{username}@cuzco.com", rol, eid_cuzco)

        # ── Empresa 2: Coquimbo ───────────────────────────────────────────────
        eid_coquimbo = _get_or_create_empresa("Coquimbo", "coquimbo")
        existing_coquimbo = set(get_all_users(eid_coquimbo)["username"].tolist()) if len(get_all_users(eid_coquimbo)) > 0 else set()
        for username, password, nombre, rol in [
            ("sup3", "Coquimbo@Sup3", "Sup 3", "admin"),
            ("sup4", "Coquimbo@Sup4", "Sup 4", "admin"),
        ]:
            if username not in existing_coquimbo:
                create_user(username, password, nombre, f"{username}@coquimbo.com", rol, eid_coquimbo)

    except Exception:
        pass


def main():
    try:
        init_db()
    except RuntimeError as e:
        st.error(f"**Error de conexion:** {e}")
        st.stop()
    _bootstrap_users()
    if "user" not in st.session_state:
        show_login()
        return

    show_sidebar()
    routes = {
        "inicio":      page_inicio,
        "cargar":      page_cargar,
        "analisis":    page_analisis,
        "historial":   page_historial,
        "estrategias": page_estrategias,
        "vicidial":    page_coquimbo if st.session_state.get("empresa_nombre") == "Coquimbo" else page_vicidial,
        "campana":     page_campana,
        "admin":       page_admin,
    }
    routes.get(st.session_state.get("page", "vicidial"), page_vicidial)()


if __name__ == "__main__":
    main()
